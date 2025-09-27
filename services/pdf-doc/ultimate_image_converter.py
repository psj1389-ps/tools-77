import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image, ImageDraw
import io
import os
import logging
import tempfile
import re

class UltimateImageConverter:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
        self.vector_graphics_found = 0

    def _robust_image_extraction(self, pdf_doc, page, img_info):
        """다양한 방법으로 이미지를 안정적으로 추출합니다 (다중화된 추출 로직)."""
        xref = img_info[0]
        
        # 방법 1: 직접 Pixmap 추출
        try:
            pix = fitz.Pixmap(pdf_doc, xref)
            if pix.width > 10 and pix.height > 10:
                self.logger.info("    - 방법 1 (직접 추출) 성공")
                return pix.tobytes("png")
        except Exception as e:
            self.logger.warning(f"    - 방법 1 실패: {e}")

        # 방법 2: 이미지 영역 렌더링 (고해상도)
        try:
            img_bbox = page.get_image_bbox(img_info)
            # 고해상도 렌더링
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=img_bbox)
            if pix.width > 30 and pix.height > 30:
                self.logger.info("    - 방법 2 (영역 렌더링) 성공")
                return pix.tobytes("png")
        except Exception as e:
            self.logger.warning(f"    - 방법 2 실패: {e}")

        # 방법 3: 이미지 영역 렌더링 (중간 해상도)
        try:
            img_bbox = page.get_image_bbox(img_info)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=img_bbox)
            if pix.width > 20 and pix.height > 20:
                self.logger.info("    - 방법 3 (중간 해상도 렌더링) 성공")
                return pix.tobytes("png")
        except Exception as e:
            self.logger.warning(f"    - 방법 3 실패: {e}")

        # 방법 4: 전체 페이지에서 이미지 영역 추출
        try:
            # 전체 페이지를 고해상도로 렌더링
            full_pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            # 이미지 위치 정보 가져오기
            img_rects = page.get_image_rects(xref)
            if img_rects:
                img_rect = img_rects[0]
                # 좌표를 해상도에 맞게 조정
                x0, y0, x1, y1 = int(img_rect.x0*2), int(img_rect.y0*2), int(img_rect.x1*2), int(img_rect.y1*2)
                
                # 영역 크기 확인
                if x1 > x0 and y1 > y0 and (x1-x0) > 20 and (y1-y0) > 20:
                    # 해당 영역만 잘라내기
                    cropped_pix = fitz.Pixmap(full_pix, fitz.IRect(x0, y0, x1, y1))
                    if cropped_pix.width > 10 and cropped_pix.height > 10:
                        self.logger.info("    - 방법 4 (페이지 영역 추출) 성공")
                        return cropped_pix.tobytes("png")
        except Exception as e:
            self.logger.warning(f"    - 방법 4 실패: {e}")

        # 방법 5: 원본 이미지 데이터 직접 추출
        try:
            img_dict = pdf_doc.extract_image(xref)
            if img_dict and 'image' in img_dict:
                img_data = img_dict['image']
                if len(img_data) > 100:  # 최소 데이터 크기 확인
                    self.logger.info("    - 방법 5 (원본 데이터 추출) 성공")
                    return img_data
        except Exception as e:
            self.logger.warning(f"    - 방법 5 실패: {e}")

        # 방법 6: 벡터 그래픽으로 변환 시도
        try:
            # 이미지가 벡터 형태일 가능성을 고려하여 해당 영역을 벡터로 렌더링
            img_rects = page.get_image_rects(xref)
            if img_rects:
                img_rect = img_rects[0]
                # 해당 영역을 고해상도로 렌더링
                pix = page.get_pixmap(matrix=fitz.Matrix(4, 4), clip=img_rect)
                if pix.width > 40 and pix.height > 40:
                    self.logger.info("    - 방법 6 (벡터 렌더링) 성공")
                    return pix.tobytes("png")
        except Exception as e:
            self.logger.warning(f"    - 방법 6 실패: {e}")

        self.logger.error("    - 모든 이미지 추출 방법 실패")
        return None

    def _detect_speech_bubbles(self, drawings):
        """말풍선 모양을 감지하는 알고리즘"""
        speech_bubbles = []
        
        for drawing in drawings:
            try:
                # 말풍선 특징 감지
                if 'items' in drawing:
                    paths = drawing['items']
                    
                    # 곡선이 포함된 경로 찾기 (말풍선의 둥근 모서리)
                    has_curves = any('c' in str(path).lower() or 'q' in str(path).lower() for path in paths)
                    
                    # 닫힌 경로인지 확인 (말풍선은 보통 닫힌 도형)
                    is_closed = any('z' in str(path).lower() for path in paths)
                    
                    # 적절한 크기인지 확인 (너무 작거나 크지 않은)
                    if 'rect' in drawing:
                        rect = drawing['rect']
                        width = abs(rect[2] - rect[0])
                        height = abs(rect[3] - rect[1])
                        
                        # 말풍선 크기 범위 (페이지 크기 대비)
                        min_size = 20
                        max_size = 300
                        
                        if (min_size < width < max_size and min_size < height < max_size and
                            has_curves and is_closed):
                            speech_bubbles.append(drawing)
                            self.logger.info(f"    - 말풍선 감지: {width}x{height}")
                            
            except Exception as e:
                self.logger.warning(f"    - 말풍선 감지 중 오류: {e}")
                continue
                
        return speech_bubbles
    
    def _detect_character_shapes(self, drawings):
        """캐릭터나 복잡한 그래픽 요소를 감지하는 알고리즘"""
        character_shapes = []
        
        for drawing in drawings:
            try:
                if 'rect' in drawing:
                    rect = drawing['rect']
                    width = abs(rect[2] - rect[0])
                    height = abs(rect[3] - rect[1])
                    
                    # 캐릭터 이미지 특징: 적당한 크기, 복잡한 경로
                    if 'items' in drawing:
                        path_count = len(drawing['items'])
                        
                        # 복잡한 경로를 가진 큰 도형 (캐릭터일 가능성)
                        if (width > 50 and height > 50 and path_count > 3):
                            character_shapes.append(drawing)
                            self.logger.info(f"    - 캐릭터 도형 감지: {width}x{height}, 경로수: {path_count}")
                            
            except Exception as e:
                self.logger.warning(f"    - 캐릭터 도형 감지 중 오류: {e}")
                continue
                
        return character_shapes

    def _extract_vector_graphics(self, page):
        """페이지에서 벡터 그래픽(선, 도형, 말풍선 등)을 추출하여 이미지로 변환합니다."""
        try:
            # 페이지의 모든 그리기 명령어 추출
            drawings = page.get_drawings()
            if not drawings:
                return None
                
            self.logger.info(f"    - {len(drawings)}개의 벡터 그래픽 발견")
            
            # 말풍선과 캐릭터 도형 감지
            speech_bubbles = self._detect_speech_bubbles(drawings)
            character_shapes = self._detect_character_shapes(drawings)
            
            self.logger.info(f"    - 말풍선 {len(speech_bubbles)}개, 캐릭터 도형 {len(character_shapes)}개 감지")
            
            # 페이지 크기 가져오기
            page_rect = page.rect
            canvas_width = int(page_rect.width * 3)  # 고해상도를 위해 3배 확대
            canvas_height = int(page_rect.height * 3)
            
            # 빈 캔버스 생성 (투명 배경)
            canvas = Image.new('RGBA', (canvas_width, canvas_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(canvas)
            
            vector_count = 0
            
            # 우선순위: 말풍선 > 캐릭터 도형 > 일반 벡터 그래픽
            priority_drawings = speech_bubbles + character_shapes
            
            for drawing in drawings:
                try:
                    # 벡터 그래픽의 경계 상자 확인
                    if 'rect' in drawing:
                        rect = drawing['rect']
                        # 좌표를 3배로 확대
                        x1, y1, x2, y2 = int(rect[0]*3), int(rect[1]*3), int(rect[2]*3), int(rect[3]*3)
                        
                        # 우선순위 그래픽인지 확인
                        is_priority = drawing in priority_drawings
                        
                        # 선 그리기 (기본적인 벡터 그래픽 표현)
                        if 'stroke' in drawing or 'fill' in drawing:
                            # 크기 확인
                            width = abs(x2-x1)
                            height = abs(y2-y1)
                            
                            if width > 15 and height > 15:  # 최소 크기 확인
                                # 우선순위 그래픽은 더 두껍게 그리기
                                line_width = 4 if is_priority else 2
                                color = (255, 0, 0, 255) if drawing in speech_bubbles else (0, 0, 255, 255) if drawing in character_shapes else (0, 0, 0, 255)
                                
                                # 복잡한 경로가 있는 경우 더 정교하게 처리
                                if 'items' in drawing and len(drawing['items']) > 1:
                                    # 복잡한 도형은 채우기로 표현
                                    draw.rectangle([x1, y1, x2, y2], fill=(*color[:3], 128), outline=color, width=line_width)
                                else:
                                    # 단순한 도형은 테두리만
                                    draw.rectangle([x1, y1, x2, y2], outline=color, width=line_width)
                                
                                vector_count += 1
                                
                except Exception as draw_error:
                    self.logger.warning(f"    - 개별 벡터 그래픽 처리 실패: {draw_error}")
                    continue
            
            if vector_count > 0:
                # 투명 배경을 흰색으로 변경
                white_bg = Image.new('RGB', (canvas_width, canvas_height), (255, 255, 255))
                white_bg.paste(canvas, mask=canvas.split()[-1] if canvas.mode == 'RGBA' else None)
                
                # PNG로 변환
                output_stream = io.BytesIO()
                white_bg.save(output_stream, format="PNG", optimize=True)
                self.vector_graphics_found += vector_count
                self.logger.info(f"    - {vector_count}개의 벡터 그래픽을 이미지로 변환 완료 (말풍선: {len(speech_bubbles)}, 캐릭터: {len(character_shapes)})")
                return output_stream.getvalue()
                
        except Exception as e:
            self.logger.warning(f"    - 벡터 그래픽 추출 실패: {e}")
            
        return None

    def _verify_and_process_image(self, img_data):
        """이미지를 검증하고 DOCX 삽입에 안전한 형식으로 처리합니다."""
        try:
            img = Image.open(io.BytesIO(img_data))

            # RGBA, P(팔레트) 등 투명도가 있는 이미지를 흰색 배경의 RGB로 변환
            if img.mode in ('RGBA', 'LA', 'P'):
                # 투명도 채널이 있는 경우에만 처리
                if 'A' in img.mode or (img.mode == 'P' and 'transparency' in img.info):
                    self.logger.info(f"    - 투명도({img.mode}) 처리 -> RGB 변환")
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    # RGBA 또는 P 모드에서 RGBA로 변환 후 처리
                    img_rgba = img.convert("RGBA")
                    background.paste(img_rgba, mask=img_rgba.split()[-1])
                    img = background

            # CMYK 이미지를 RGB로 변환
            elif img.mode == 'CMYK':
                self.logger.info("    - CMYK 처리 -> RGB 변환")
                img = img.convert("RGB")

            # 최종적으로 PNG 형식으로 메모리에 저장
            output_stream = io.BytesIO()
            img.save(output_stream, format="PNG", optimize=True)
            return output_stream.getvalue()
        except Exception as e:
            self.logger.error(f"    - 이미지 처리 실패: {e}")
            return None

    def _extract_embedded_images_alternative(self, page):
        """페이지에서 임베디드 이미지를 대체 방법으로 추출합니다."""
        extracted_images = []
        
        try:
            # 페이지의 모든 객체 검사
            for obj in page.get_contents():
                try:
                    # 이미지 객체 찾기
                    if hasattr(obj, 'get_pixmap'):
                        pix = obj.get_pixmap()
                        if pix and pix.width > 10 and pix.height > 10:
                            extracted_images.append(pix.tobytes("png"))
                            self.logger.info(f"    - 대체 방법으로 이미지 추출 성공: {pix.width}x{pix.height}")
                except Exception:
                    continue
                    
        except Exception as e:
            self.logger.warning(f"    - 대체 이미지 추출 실패: {e}")
            
        return extracted_images
    
    def _merge_vector_and_raster_images(self, page, raster_images, vector_image):
        """래스터 이미지와 벡터 그래픽을 통합하여 처리합니다."""
        merged_images = []
        
        try:
            # 래스터 이미지들 추가
            for img_data in raster_images:
                if img_data:
                    merged_images.append({
                        'type': 'raster',
                        'data': img_data,
                        'priority': 1
                    })
            
            # 벡터 그래픽 추가
            if vector_image:
                merged_images.append({
                    'type': 'vector',
                    'data': vector_image,
                    'priority': 2
                })
            
            # 우선순위에 따라 정렬 (래스터 이미지 우선)
            merged_images.sort(key=lambda x: x['priority'])
            
            self.logger.info(f"    - 통합된 이미지: 래스터 {len(raster_images)}개, 벡터 {1 if vector_image else 0}개")
            
        except Exception as e:
            self.logger.warning(f"    - 이미지 통합 처리 실패: {e}")
            
        return merged_images

    def _add_page_to_docx_image_priority(self, doc, text_blocks, merged_images):
        """이미지 우선 모드: 이미지를 먼저 배치하고 텍스트를 배치합니다."""
        try:
            # 이미지들을 먼저 추가
            for img_info in merged_images:
                if img_info['type'] == 'raster':
                    self._insert_image_to_docx(doc, img_info['data'])
                elif img_info['type'] == 'vector':
                    self._insert_image_to_docx(doc, img_info['data'])
            
            # 텍스트 블록들을 y 좌표 순으로 정렬하여 추가
            sorted_blocks = sorted(text_blocks, key=lambda x: x['y'])
            for block in sorted_blocks:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block['text'])
                run.font.size = Pt(block.get('font_size', 11))
                
            self.logger.info("    ✅ 이미지 우선 모드로 페이지 추가 완료")
            
        except Exception as e:
            self.logger.error(f"    ❌ 이미지 우선 모드 페이지 추가 실패: {e}")
    
    def _add_page_to_docx_text_priority(self, doc, text_blocks, merged_images):
        """텍스트 우선 모드: 텍스트를 먼저 배치하고 이미지를 배치합니다."""
        try:
            # 텍스트 블록들을 먼저 추가
            sorted_blocks = sorted(text_blocks, key=lambda x: x['y'])
            for block in sorted_blocks:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(block['text'])
                run.font.size = Pt(block.get('font_size', 11))
            
            # 이미지들을 나중에 추가
            for img_info in merged_images:
                if img_info['type'] == 'raster':
                    self._insert_image_to_docx(doc, img_info['data'])
                elif img_info['type'] == 'vector':
                    self._insert_image_to_docx(doc, img_info['data'])
                    
            self.logger.info("    ✅ 텍스트 우선 모드로 페이지 추가 완료")
            
        except Exception as e:
            self.logger.error(f"    ❌ 텍스트 우선 모드 페이지 추가 실패: {e}")
    
    def _insert_image_to_docx(self, doc, img_data):
        """DOCX 문서에 이미지를 삽입합니다."""
        try:
            # 임시 파일로 이미지 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(img_data)
                temp_path = temp_file.name
            
            # DOCX에 이미지 추가
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            
            # 이미지 크기 조정
            try:
                from PIL import Image
                with Image.open(temp_path) as img:
                    width, height = img.size
                    # 최대 크기 제한 (A4 용지 기준)
                    max_width = Inches(6)
                    max_height = Inches(8)
                    
                    if width > height:
                        new_width = min(max_width, Inches(width/100))
                        run.add_picture(temp_path, width=new_width)
                    else:
                        new_height = min(max_height, Inches(height/100))
                        run.add_picture(temp_path, height=new_height)
            except Exception:
                # PIL이 없거나 이미지 처리 실패 시 기본 크기로 삽입
                run.add_picture(temp_path, width=Inches(4))
            
            # 임시 파일 삭제
            os.unlink(temp_path)
            
        except Exception as e:
             self.logger.warning(f"    - 이미지 삽입 실패: {e}")

    def _analyze_presentation_layout(self, page, text_blocks, images):
        """프레젠테이션 형태 문서의 레이아웃을 분석하여 캐릭터와 텍스트의 연관성을 파악합니다."""
        layout_analysis = {
            'character_regions': [],
            'speech_bubble_regions': [],
            'text_image_associations': [],
            'layout_type': 'unknown'
        }
        
        try:
            page_width = page.rect.width
            page_height = page.rect.height
            
            # 이미지 영역 분석
            for i, img_info in enumerate(images):
                try:
                    img_rects = page.get_image_rects(img_info[0])
                    if img_rects:
                        img_rect = img_rects[0]
                        
                        # 이미지 크기와 위치로 캐릭터 이미지 추정
                        img_width = img_rect.width
                        img_height = img_rect.height
                        img_area = img_width * img_height
                        
                        # 캐릭터 이미지 특성 분석
                        is_character = self._is_likely_character_image(img_rect, page_width, page_height)
                        is_speech_bubble = self._is_likely_speech_bubble(img_rect, page_width, page_height)
                        
                        if is_character:
                            layout_analysis['character_regions'].append({
                                'index': i,
                                'rect': img_rect,
                                'area': img_area,
                                'position': 'left' if img_rect.x0 < page_width/2 else 'right'
                            })
                        
                        if is_speech_bubble:
                            layout_analysis['speech_bubble_regions'].append({
                                'index': i,
                                'rect': img_rect,
                                'area': img_area
                            })
                            
                except Exception as e:
                    self.logger.warning(f"    - 이미지 {i} 분석 실패: {e}")
            
            # 텍스트와 이미지의 연관성 분석
            for text_block in text_blocks:
                text_rect = fitz.Rect(text_block['x'], text_block['y'], 
                                    text_block['x'] + text_block['width'], 
                                    text_block['y'] + text_block['height'])
                
                # 가장 가까운 캐릭터 이미지 찾기
                closest_character = None
                min_distance = float('inf')
                
                for char_region in layout_analysis['character_regions']:
                    distance = self._calculate_distance(text_rect, char_region['rect'])
                    if distance < min_distance:
                        min_distance = distance
                        closest_character = char_region
                
                if closest_character and min_distance < page_width * 0.3:  # 페이지 너비의 30% 이내
                    layout_analysis['text_image_associations'].append({
                        'text_block': text_block,
                        'character_region': closest_character,
                        'distance': min_distance,
                        'relationship': 'dialogue' if min_distance < page_width * 0.15 else 'related'
                    })
            
            # 레이아웃 타입 결정
            if len(layout_analysis['character_regions']) > 0 and len(layout_analysis['speech_bubble_regions']) > 0:
                layout_analysis['layout_type'] = 'comic_presentation'
            elif len(layout_analysis['character_regions']) > 0:
                layout_analysis['layout_type'] = 'character_presentation'
            elif len(layout_analysis['text_image_associations']) > 0:
                layout_analysis['layout_type'] = 'mixed_presentation'
            else:
                layout_analysis['layout_type'] = 'text_presentation'
            
            self.logger.info(f"    📋 레이아웃 분석 완료: {layout_analysis['layout_type']}")
            self.logger.info(f"      - 캐릭터 영역: {len(layout_analysis['character_regions'])}개")
            self.logger.info(f"      - 말풍선 영역: {len(layout_analysis['speech_bubble_regions'])}개")
            self.logger.info(f"      - 텍스트-이미지 연관: {len(layout_analysis['text_image_associations'])}개")
            
        except Exception as e:
            self.logger.error(f"    ❌ 프레젠테이션 레이아웃 분석 실패: {e}")
            
        return layout_analysis
    
    def _is_likely_character_image(self, img_rect, page_width, page_height):
        """이미지가 캐릭터 이미지일 가능성을 판단합니다."""
        try:
            img_width = img_rect.width
            img_height = img_rect.height
            
            # 캐릭터 이미지 특성
            # 1. 세로가 가로보다 길거나 비슷함 (인물 특성)
            aspect_ratio = img_height / img_width if img_width > 0 else 0
            
            # 2. 적절한 크기 (너무 작지도 크지도 않음)
            size_ratio = (img_width * img_height) / (page_width * page_height)
            
            # 3. 페이지 가장자리에 위치 (캐릭터는 보통 좌우 가장자리)
            is_edge_positioned = (img_rect.x0 < page_width * 0.3) or (img_rect.x1 > page_width * 0.7)
            
            return (aspect_ratio >= 0.8 and aspect_ratio <= 2.5 and 
                   size_ratio >= 0.05 and size_ratio <= 0.4 and 
                   is_edge_positioned)
                   
        except Exception:
            return False
    
    def _is_likely_speech_bubble(self, img_rect, page_width, page_height):
        """이미지가 말풍선일 가능성을 판단합니다."""
        try:
            img_width = img_rect.width
            img_height = img_rect.height
            
            # 말풍선 특성
            # 1. 가로가 세로보다 길거나 비슷함 (말풍선 특성)
            aspect_ratio = img_width / img_height if img_height > 0 else 0
            
            # 2. 중간 크기 (텍스트를 담을 수 있는 크기)
            size_ratio = (img_width * img_height) / (page_width * page_height)
            
            # 3. 페이지 중앙 부근에 위치
            is_center_positioned = (img_rect.x0 > page_width * 0.2 and img_rect.x1 < page_width * 0.8)
            
            return (aspect_ratio >= 0.5 and aspect_ratio <= 3.0 and 
                   size_ratio >= 0.02 and size_ratio <= 0.25 and 
                   is_center_positioned)
                   
        except Exception:
            return False
    
    def _calculate_distance(self, rect1, rect2):
        """두 사각형 간의 거리를 계산합니다."""
        try:
            # 중심점 간의 거리 계산
            center1_x = (rect1.x0 + rect1.x1) / 2
            center1_y = (rect1.y0 + rect1.y1) / 2
            center2_x = (rect2.x0 + rect2.x1) / 2
            center2_y = (rect2.y0 + rect2.y1) / 2
            
            return ((center1_x - center2_x) ** 2 + (center1_y - center2_y) ** 2) ** 0.5
        except Exception:
            return float('inf')

    def convert_with_guaranteed_images(self, pdf_path, output_path, mode='balanced'):
        """이미지 누락을 방지하고 원본 레이아웃을 최대한 보존하는 PDF 변환 메서드 (다중화된 추출 시스템)"""
        # 변환 모드 설정
        extraction_modes = {
            'image': {'priority': 'image_first', 'quality': 'ultra'},
            'balanced': {'priority': 'balanced', 'quality': 'high'},
            'text': {'priority': 'text_first', 'quality': 'medium'}
        }
        
        current_mode = extraction_modes.get(mode, extraction_modes['balanced'])
        self.logger.info(f"🚀 변환 모드: {mode} (우선순위: {current_mode['priority']}, 품질: {current_mode['quality']})")
        self.logger.info(f"🚀 '{pdf_path}' 변환 시작 (레이아웃 보존 + 이미지 보존 최우선 모드)")
        
        # 이미지 추출 통계
        extraction_stats = {
            'total_images': 0,
            'successful_extractions': 0,
            'vector_graphics': 0,
            'failed_extractions': 0
        }
        try:
            pdf_doc = fitz.open(pdf_path)
            docx_doc = Document()
            
            # 첫 번째 페이지로 문서 방향 설정
            if len(pdf_doc) > 0:
                first_page = pdf_doc.load_page(0)
                page_rect = first_page.rect
                page_width = page_rect.width
                page_height = page_rect.height
                
                # 문서 방향 결정 (가로 > 세로면 가로형, 아니면 세로형)
                section = docx_doc.sections[0]
                if page_width > page_height:
                    # 가로형 문서
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11)
                    section.page_height = Inches(8.5)
                    self.logger.info("📐 문서 방향: 가로형 (Landscape)")
                else:
                    # 세로형 문서
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Inches(8.5)
                    section.page_height = Inches(11)
                    self.logger.info("📐 문서 방향: 세로형 (Portrait)")
                
                # 여백 설정을 원본과 유사하게 조정
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
            
            images_added = 0
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                self.logger.info(f"\n📄 페이지 {page_num + 1} 처리 중...")

                # 페이지 내 요소들의 위치 정보를 수집하여 순서대로 배치
                page_elements = []
                
                # 텍스트 블록들의 위치 정보 수집
                text_blocks = page.get_text("dict")
                text_blocks_info = []
                for block in text_blocks.get("blocks", []):
                    if "lines" in block:
                        block_text = ""
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                block_text += span.get("text", "")
                        if block_text.strip():
                            text_block_info = {
                                "text": block_text.strip(),
                                "x": block["bbox"][0],
                                "y": block["bbox"][1],
                                "width": block["bbox"][2] - block["bbox"][0],
                                "height": block["bbox"][3] - block["bbox"][1]
                            }
                            text_blocks_info.append(text_block_info)
                            page_elements.append({
                                "type": "text",
                                "content": block_text.strip(),
                                "y_position": block["bbox"][1],  # 상단 y 좌표
                                "bbox": block["bbox"]
                            })
                
                # 프레젠테이션 레이아웃 분석 (캐릭터와 텍스트 연관성 분석)
                layout_analysis = self._analyze_presentation_layout(page, text_blocks_info, image_list)
                
                # 이미지들의 위치 정보 수집
                image_list = page.get_images(full=True)
                for img_index, img_info in enumerate(image_list):
                    # 이미지의 위치 정보 가져오기
                    img_rects = page.get_image_rects(img_info[0])
                    if img_rects:
                        img_rect = img_rects[0]  # 첫 번째 위치 사용
                        page_elements.append({
                            "type": "image",
                            "content": img_info,
                            "y_position": img_rect.y0,  # 상단 y 좌표
                            "bbox": (img_rect.x0, img_rect.y0, img_rect.x1, img_rect.y1),
                            "index": img_index
                        })
                
                # 벡터 그래픽도 위치 정보와 함께 처리
                vector_img_data = self._extract_vector_graphics(page)
                if vector_img_data:
                    # 벡터 그래픽을 적절한 위치에 배치
                    page_elements.append({
                        "type": "vector",
                        "content": vector_img_data,
                        "y_position": page.rect.height * 0.1,  # 페이지 상단 근처
                        "bbox": (0, 0, page.rect.width, page.rect.height * 0.2)
                    })
                
                # y 좌표 기준으로 정렬 (위에서 아래로)
                page_elements.sort(key=lambda x: x["y_position"])
                
                # 정렬된 순서대로 요소들을 DOCX에 추가
                for element in page_elements:
                    if element["type"] == "text":
                        # 텍스트 추가 (원본 서식 유지 시도)
                        paragraph = docx_doc.add_paragraph(element["content"])
                        # 텍스트 위치에 따른 정렬 설정
                        bbox = element["bbox"]
                        page_width = page.rect.width
                        text_center_x = (bbox[0] + bbox[2]) / 2
                        
                        if text_center_x < page_width * 0.25:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif text_center_x > page_width * 0.75:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    elif element["type"] == "image":
                        # 이미지 처리 및 추가
                        img_info = element["content"]
                        img_index = element["index"]
                        self.logger.info(f"  - 이미지 {img_index + 1} 처리 시작 (위치 기반 배치)...")
                        
                        # 강력한 이미지 추출
                        raw_img_data = self._robust_image_extraction(pdf_doc, page, img_info)
                        if not raw_img_data:
                            continue

                        # 안전한 이미지 처리
                        processed_img_data = self._verify_and_process_image(raw_img_data)
                        if not processed_img_data:
                            continue

                        # DOCX에 이미지 삽입
                        try:
                            paragraph = docx_doc.add_paragraph()
                            
                            # 이미지 위치에 따른 정렬 설정
                            bbox = element["bbox"]
                            page_width = page.rect.width
                            img_center_x = (bbox[0] + bbox[2]) / 2
                            
                            if img_center_x < page_width * 0.3:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif img_center_x > page_width * 0.7:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            run = paragraph.add_run()
                            
                            # 임시 파일로 이미지 저장
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                                temp_file.write(processed_img_data)
                                temp_filepath = temp_file.name
                            
                            # 원본 이미지 크기 정보 활용
                            img = Image.open(io.BytesIO(processed_img_data))
                            aspect_ratio = img.width / img.height
                            
                            # 원본 PDF에서의 이미지 크기 비율 계산
                            pdf_img_width = bbox[2] - bbox[0]
                            pdf_img_height = bbox[3] - bbox[1]
                            page_width_ratio = pdf_img_width / page.rect.width
                            page_height_ratio = pdf_img_height / page.rect.height
                            
                            # 페이지 방향에 따른 최대 이미지 크기 설정
                            section = docx_doc.sections[0]
                            if section.orientation == WD_ORIENT.LANDSCAPE:
                                max_width = Inches(10)
                                max_height = Inches(7)
                            else:
                                max_width = Inches(7)
                                max_height = Inches(9)
                            
                            # 원본 비율을 더 정확하게 반영한 크기 설정
                            if page_width_ratio > 0.9:  # 페이지 너비의 90% 이상 (전체 너비)
                                img_width = max_width * 0.95
                            elif page_width_ratio > 0.7:  # 페이지 너비의 70% 이상 (큰 이미지)
                                img_width = max_width * 0.8
                            elif page_width_ratio > 0.4:  # 페이지 너비의 40% 이상 (중간 이미지)
                                img_width = max_width * 0.6
                            elif page_width_ratio > 0.2:  # 페이지 너비의 20% 이상 (작은 이미지)
                                img_width = max_width * 0.4
                            else:  # 매우 작은 이미지
                                img_width = max_width * 0.25
                            
                            # 높이 비율도 고려하여 조정
                            if page_height_ratio > 0.6:  # 세로로 긴 이미지
                                img_height = min(max_height * 0.8, img_width / aspect_ratio)
                                img_width = img_height * aspect_ratio
                            
                            # 최소/최대 크기 제한
                            img_width = max(Inches(0.8), min(img_width, max_width))
                            
                            run.add_picture(temp_filepath, width=img_width)
                            
                            # 임시 파일 정리
                            try:
                                os.remove(temp_filepath)
                            except Exception:
                                pass
                            
                            images_added += 1
                            self.logger.info(f"    ✅ 이미지 {img_index + 1} 삽입 성공 (위치 기반 배치)!")
                        except Exception as e:
                            self.logger.error(f"    ❌ 이미지 {img_index + 1} 삽입 실패: {e}")
                            try:
                                if 'temp_filepath' in locals() and os.path.exists(temp_filepath):
                                    os.remove(temp_filepath)
                            except Exception:
                                pass

                    elif element["type"] == "vector":
                        # 벡터 그래픽 처리
                        try:
                            paragraph = docx_doc.add_paragraph()
                            
                            # 벡터 그래픽 위치에 따른 정렬 설정
                            bbox = element["bbox"]
                            page_width = page.rect.width
                            vector_center_x = (bbox[0] + bbox[2]) / 2
                            
                            if vector_center_x < page_width * 0.25:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif vector_center_x > page_width * 0.75:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            run = paragraph.add_run()
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                                temp_file.write(element["content"])
                                temp_filepath = temp_file.name
                            
                            # 벡터 그래픽 크기를 원본 비율에 맞게 조정
                            vector_width_ratio = (bbox[2] - bbox[0]) / page.rect.width
                            section = docx_doc.sections[0]
                            if section.orientation == WD_ORIENT.LANDSCAPE:
                                max_vector_width = Inches(8)
                            else:
                                max_vector_width = Inches(6)
                            
                            vector_width = max_vector_width * min(vector_width_ratio * 1.2, 0.9)
                            vector_width = max(Inches(2), min(vector_width, max_vector_width))
                            
                            run.add_picture(temp_filepath, width=vector_width)
                            
                            try:
                                os.remove(temp_filepath)
                            except Exception:
                                pass
                                
                            self.logger.info("    ✅ 벡터 그래픽 삽입 성공 (위치 기반 배치)!")
                        except Exception as e:
                            self.logger.error(f"    ❌ 벡터 그래픽 삽입 실패: {e}")
                
                # 페이지 요소 통계 출력
                img_count = len([e for e in page_elements if e['type'] == 'image'])
                text_count = len([e for e in page_elements if e['type'] == 'text'])
                vector_count = len([e for e in page_elements if e['type'] == 'vector'])
                self.logger.info(f"  - 총 {img_count}개의 이미지, {text_count}개의 텍스트 블록, {vector_count}개의 벡터 그래픽을 위치 기반으로 배치")

                if page_num < len(pdf_doc) - 1:
                    docx_doc.add_page_break()

            pdf_doc.close()
            docx_doc.save(output_path)
            
            # 변환 통계 출력
            success_rate = (extraction_stats['successful_extractions'] / max(extraction_stats['total_images'], 1)) * 100
            self.logger.info(f"\n📊 이미지 추출 통계:")
            self.logger.info(f"  - 총 이미지: {extraction_stats['total_images']}개")
            self.logger.info(f"  - 성공적 추출: {extraction_stats['successful_extractions']}개")
            self.logger.info(f"  - 벡터 그래픽: {extraction_stats['vector_graphics']}개")
            self.logger.info(f"  - 실패: {extraction_stats['failed_extractions']}개")
            self.logger.info(f"  - 성공률: {success_rate:.1f}%")
            
            self.logger.info(f"\n🎉 변환 완료! 총 {images_added}개의 이미지와 {self.vector_graphics_found}개의 벡터 그래픽이 성공적으로 삽입되었습니다.")
            return True

        except Exception as e:
            self.logger.error(f"전체 변환 과정에서 오류 발생: {e}")
            return False

# ----- 사용 예시 -----
if __name__ == '__main__':
    converter = UltimateImageConverter()
    
    # 기본 테스트 파일을 test_document.pdf로 설정
    pdf_file = "test_document.pdf"
    output_docx = "final_output_with_images.docx"
    
    if os.path.exists(pdf_file):
        converter.convert_with_guaranteed_images(pdf_file, output_docx)
    else:
        print(f"'{pdf_file}' 파일을 찾을 수 없습니다.")