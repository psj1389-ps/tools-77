from flask import Flask, request, render_template, send_file, jsonify
from dotenv import load_dotenv
import os
import tempfile
import subprocess
import platform
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Mm
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import pytesseract
import cv2
import numpy as np
import json
import logging
import zipfile

# Adobe SDK 임포트 - 선택적 로딩 (SDK 4.2 구조)
try:
    # 올바른 Credentials 클래스 import (SDK 4.2)
    from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
    from adobe.pdfservices.operation.pdf_services import PDFServices
    from adobe.pdfservices.operation.pdfjobs.jobs.extract_pdf_job import ExtractPDFJob
    from adobe.pdfservices.operation.pdfjobs.params.extract_pdf.extract_pdf_params import ExtractPDFParams
    from adobe.pdfservices.operation.pdfjobs.params.extract_pdf.extract_element_type import ExtractElementType
    from adobe.pdfservices.operation.pdfjobs.result.extract_pdf_result import ExtractPDFResult
    from adobe.pdfservices.operation.io.stream_asset import StreamAsset
    from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
    
    # ExportPDFOperation 관련 import (SDK 4.2)
    from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
    from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult
    try:
        from adobe.pdfservices.operation.io.media_type import MediaType
    except ImportError:
        # SDK 버전에 따라 MediaType 경로가 다를 수 있음
        try:
            from adobe.pdfservices.operation.io.asset import MediaType
        except ImportError:
            # MediaType을 직접 정의
            class MediaType:
                PDF = "application/pdf"
    # Exception 클래스들 (SDK 4.2 올바른 경로)
    try:
        from adobe.pdfservices.operation.exception.exceptions import ServiceApiException
        from adobe.pdfservices.operation.exception.exceptions import ServiceUsageException
    except ImportError:
        # SDK 버전에 따라 exception 경로가 다를 수 있음
        ServiceApiException = Exception
        ServiceUsageException = Exception
    ADOBE_SDK_AVAILABLE = True
    print("✅ Adobe PDF Services SDK 4.2.0 로드 완료")
except ImportError as e:
    ADOBE_SDK_AVAILABLE = False
    ServiceApiException = Exception
    ServiceUsageException = Exception
    print(f"ℹ️ Adobe PDF Services SDK 미설치 - 고급 OCR 모드로 동작: {e}")

load_dotenv()
app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

# Vite 클라이언트 및 정적 파일 서빙을 위한 라우트 추가
@app.route('/@vite/client')
def vite_client():
    """Vite 클라이언트 요청 처리"""
    return '', 404  # 개발 모드가 아니므로 404 반환

@app.route('/assets/<path:filename>')
def serve_assets(filename):
    """정적 자산 파일 서빙"""
    return send_file(os.path.join('static', 'assets', filename))

# Adobe PDF Services API 설정 - OAuth Server-to-Server 방식
ADOBE_CLIENT_ID = os.getenv("ADOBE_CLIENT_ID")
ADOBE_CLIENT_SECRET = os.getenv("ADOBE_CLIENT_SECRET")
ADOBE_ORGANIZATION_ID = os.getenv("ADOBE_ORGANIZATION_ID")

# 폴더 생성
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename, content_type=None):
    """파일 형식 확인 (확장자 또는 MIME 타입 기반)"""
    # 확장자로 확인
    if '.' in filename:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in {'pdf', 'docx'}
    
    # MIME 타입으로 확인 (확장자가 없는 경우)
    if content_type:
        return ('pdf' in content_type or 
                'document' in content_type or 
                'word' in content_type)
    
    return False

def setup_korean_fonts():
    """한글 폰트 설정"""
    try:
        # Windows 시스템 폰트 경로들
        font_paths = [
            r"C:\Windows\Fonts\malgun.ttf",  # 맑은 고딕
            r"C:\Windows\Fonts\gulim.ttc",   # 굴림
            r"C:\Windows\Fonts\batang.ttc",  # 바탕
            r"C:\Windows\Fonts\dotum.ttc",   # 돋움
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('Korean', font_path))
                    print(f"한글 폰트 등록 성공: {font_path}")
                    return True
                except Exception as e:
                    print(f"폰트 등록 실패: {font_path} - {e}")
                    continue
        
        print("한글 폰트를 찾을 수 없습니다. 기본 폰트를 사용합니다.")
        return False
    except Exception as e:
        print(f"폰트 설정 오류: {e}")
        return False

def setup_korean_font(doc):
    """Word 문서용 한글 폰트 설정"""
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 문서 스타일 설정
        styles = doc.styles
        
        # Normal 스타일에 한글 폰트 설정
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = '맑은 고딕'
        font.size = Pt(11)
        
        # 동아시아 폰트 설정
        font.element.set(qn('w:eastAsia'), '맑은 고딕')
        
        print("✅ Word 문서 한글 폰트 설정 완료: 맑은 고딕")
        return True
        
    except Exception as e:
        print(f"⚠️ Word 문서 한글 폰트 설정 실패: {e}")
        return False

def convert_pdf_to_docx_with_adobe(pdf_path, output_path):
    """Adobe PDF Services SDK 4.2를 사용하여 PDF를 DOCX로 직접 변환 (ExportPDFJob 사용)"""
    if not ADOBE_SDK_AVAILABLE:
        print("⚠️ Adobe SDK가 설치되지 않아 False 반환")
        return False
        
    try:
        print(f"🔗 Adobe ExportPDF API 호출: {pdf_path} -> {output_path}")
        
        # 파일 존재 확인
        if not os.path.exists(pdf_path):
            print(f"❌ PDF 파일이 존재하지 않습니다: {pdf_path}")
            return False
            
        # 파일 크기 확인 (100MB 제한)
        file_size = os.path.getsize(pdf_path)
        if file_size > 100 * 1024 * 1024:
            print(f"❌ 파일 크기가 너무 큽니다: {file_size / (1024*1024):.1f}MB")
            return False
        
        # 1. Adobe API 자격 증명 설정 (SDK 4.2)
        client_id = ADOBE_CLIENT_ID
        client_secret = ADOBE_CLIENT_SECRET
        
        if not client_id or not client_secret:
            print("❌ Adobe 자격 증명이 설정되지 않았습니다")
            return False
        
        # ServicePrincipalCredentials 사용 (SDK 4.2 올바른 방식)
        credentials = ServicePrincipalCredentials(
            client_id=client_id,
            client_secret=client_secret
        )
        
        # 2. PDF Services 클라이언트 생성
        pdf_services = PDFServices(credentials=credentials)
        
        # 3. PDF 파일을 업로드하여 Asset 생성 (SDK 4.2 권장 방식)
        with open(pdf_path, 'rb') as file:
            input_stream = file.read()
        input_asset = pdf_services.upload(input_stream=input_stream, mime_type='application/pdf')
        
        # 4. Export 파라미터 설정: DOCX 포맷으로 지정
        export_pdf_params = ExportPDFParams(target_format=ExportPDFTargetFormat.DOCX)
        
        # 5. ExportPDFJob 생성
        export_pdf_job = ExportPDFJob(
            input_asset=input_asset,
            export_pdf_params=export_pdf_params
        )
        
        print(f"📤 Adobe SDK 4.2로 PDF->DOCX 변환 중... ({file_size / 1024:.1f}KB)")
        
        print(">>> [DEBUG 1] Adobe 변환 함수 진입")
        try:
            print(">>> [DEBUG 2] try 블록 진입, execute() 호출 직전")
            
            # 6. 작업 제출 및 결과 대기 - 실제 Adobe API 실행 지점
            location = pdf_services.submit(export_pdf_job)
            pdf_services_response = pdf_services.get_job_result(location, ExportPDFResult)
            
            print(">>> [DEBUG 3] execute() 호출 성공")
            conversion_success = True  # 성공했음을 표시
            
        except ServiceApiException as e:
            # Adobe API 관련 에러 (가장 흔함)
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            print(f"❌ Adobe ServiceApiException 발생: {e}")
            print(f"    - Status Code: {e.status_code if hasattr(e, 'status_code') else 'N/A'}")
            print(f"    - Error Code: {e.error_code if hasattr(e, 'error_code') else 'N/A'}")
            print(f"    - Error Message: {e.message if hasattr(e, 'message') else str(e)}")
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            conversion_success = False
            raise  # 기존 예외 처리로 전달
            
        except Exception as e:
            # 그 외 모든 예상치 못한 에러
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            print(f"❌ 변환 중 알 수 없는 예외 발생: {str(e)}")
            import traceback
            traceback.print_exc()
            print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            conversion_success = False
            raise  # 기존 예외 처리로 전달
            
        print(">>> [DEBUG 4] Adobe 변환 함수 종료")
        
        # 7. 결과 확인 및 저장 (CloudAsset 오류 해결)
        result = pdf_services_response.get_result()
        result_asset = result.get_asset()
        
        # 결과를 파일로 저장 - CloudAsset 타입별 처리
        download_success = False
        
        # StreamAsset 처리 (우선 시도)
        try:
            if hasattr(result_asset, 'get_input_stream'):
                input_stream = result_asset.get_input_stream()
                
                # 스트림 데이터를 bytes로 변환
                if hasattr(input_stream, 'read'):
                    stream_data = input_stream.read()
                elif hasattr(input_stream, '__enter__'):
                    with input_stream as stream:
                        stream_data = stream.read()
                else:
                    stream_data = input_stream
                
                # bytes 타입 확인 및 저장
                if isinstance(stream_data, bytes):
                    with open(output_path, 'wb') as output_file:
                        output_file.write(stream_data)
                    download_success = True
                else:
                    # str이나 다른 타입인 경우 bytes로 변환
                    if isinstance(stream_data, str):
                        stream_data = stream_data.encode('utf-8')
                    elif hasattr(stream_data, 'encode'):
                        stream_data = stream_data.encode()
                    
                    with open(output_path, 'wb') as output_file:
                        output_file.write(stream_data)
                    download_success = True
        except Exception as stream_error:
            print(f"⚠️ StreamAsset 처리 실패: {stream_error}")
        
        # CloudAsset 처리 (대안 방법)
        if not download_success:
            try:
                if hasattr(result_asset, 'download_uri'):
                    import requests
                    response = requests.get(result_asset.download_uri)
                    response.raise_for_status()
                    
                    with open(output_path, 'wb') as output_file:
                        output_file.write(response.content)
                    download_success = True
                elif hasattr(result_asset, 'get_stream'):
                    stream_data = result_asset.get_stream()
                    if hasattr(stream_data, 'read'):
                        raw_data = stream_data.read()
                    else:
                        raw_data = stream_data
                    
                    if isinstance(raw_data, bytes):
                        final_data = raw_data
                    elif isinstance(raw_data, str):
                        final_data = raw_data.encode('utf-8')
                    else:
                        final_data = str(raw_data).encode('utf-8')
                    
                    with open(output_path, 'wb') as output_file:
                        output_file.write(final_data)
                    download_success = True
                else:
                    # 직접 bytes 변환 시도
                    if isinstance(result_asset, bytes):
                        with open(output_path, 'wb') as output_file:
                            output_file.write(result_asset)
                        download_success = True
                    else:
                        raise Exception(f"지원되지 않는 Asset 타입: {type(result_asset)}")
            except Exception as alt_error:
                print(f"⚠️ CloudAsset 처리 실패: {alt_error}")
        
        if not download_success:
            raise Exception(f"결과 저장 실패 - Asset 타입: {type(result_asset)}, 속성: {dir(result_asset)}")
        
        print(f"✅ Adobe SDK 4.2 변환 성공: {pdf_path} -> {output_path} (편집 가능한 DOCX)")
        return True
        
    except ServiceApiException as api_error:
        print(f"❌ Adobe ExportPDF ServiceApiException 발생:")
        print(f"   - 에러 메시지: {api_error}")
        print(f"   - 에러 타입: {type(api_error)}")
        if hasattr(api_error, 'status_code'):
            print(f"   - HTTP 상태 코드: {api_error.status_code}")
        if hasattr(api_error, 'error_code'):
            print(f"   - Adobe 에러 코드: {api_error.error_code}")
        if hasattr(api_error, 'message'):
            print(f"   - 상세 메시지: {api_error.message}")
        
        # HTTP 400 에러 특별 처리
        if hasattr(api_error, 'status_code') and api_error.status_code == 400:
            print("💡 HTTP 400 Bad Request - 요청 파라미터나 파일 형식을 확인해주세요")
            print("   - PDF 파일이 손상되었거나 지원되지 않는 형식일 수 있습니다")
            print("   - Adobe API 요청 파라미터를 확인해주세요")
        
        return False
        
    except Exception as e:
        error_msg = str(e)
        print(f"❌ Adobe ExportPDF 일반 오류: {error_msg}")
        print(f"   - 에러 타입: {type(e)}")
        
        # 구체적인 오류 메시지 분석
        if "credentials" in error_msg.lower() or "authentication" in error_msg.lower():
            print("💡 인증 오류 - Adobe API 키를 확인해주세요")
        elif "timeout" in error_msg.lower() or "connection" in error_msg.lower():
            print("💡 네트워크 오류 - 인터넷 연결 또는 Adobe 서버 상태를 확인해주세요")
        elif "file" in error_msg.lower() and "corrupt" in error_msg.lower():
            print("💡 파일 손상 오류 - PDF 파일이 손상되었을 수 있습니다")
        elif "400" in error_msg:
            print("💡 HTTP 400 에러 - 요청이 잘못되었거나 파일이 손상되었을 수 있습니다")
        
        return False

def extract_with_adobe(pdf_path):
    """Adobe PDF Services Extract API를 사용하여 텍스트와 좌표 정보 추출 (SDK 4.2) - 우선 사용"""
    if not ADOBE_SDK_AVAILABLE:
        print("⚠️ Adobe SDK가 설치되지 않아 None 반환")
        return None
        
    try:
        print(f"🔗 Adobe Extract API 호출: {pdf_path}")
        
        # 파일 존재 확인
        if not os.path.exists(pdf_path):
            print(f"❌ PDF 파일이 존재하지 않습니다: {pdf_path}")
            return None
            
        # 파일 크기 확인 (100MB 제한)
        file_size = os.path.getsize(pdf_path)
        if file_size > 100 * 1024 * 1024:
            print(f"❌ 파일 크기가 너무 큽니다: {file_size / (1024*1024):.1f}MB")
            return None
        
        # 자격 증명 설정 (SDK 4.2 호환) - 환경변수 사용
        client_id = ADOBE_CLIENT_ID
        client_secret = ADOBE_CLIENT_SECRET
        
        if not client_id or not client_secret:
            print("❌ Adobe 자격 증명이 설정되지 않았습니다")
            return None
        
        # ServicePrincipalCredentials 사용 (SDK 4.2 정확한 클래스)
        credentials = ServicePrincipalCredentials(
            client_id=client_id,
            client_secret=client_secret
        )
        
        # PDF Services 클라이언트 생성
        pdf_services = PDFServices(credentials=credentials)
        
        print(f"📤 Adobe SDK 4.2로 파일 처리 중... ({file_size / 1024:.1f}KB)")
        
        # Extract 파라미터 설정 - 텍스트와 테이블 추출 (FIGURES 오류 방지)
        extract_pdf_params = ExtractPDFParams(
            elements_to_extract=[
                ExtractElementType.TEXT,    # 텍스트 추출
                ExtractElementType.TABLES   # 테이블 추출
            ]
        )
        
        # PDF 파일을 업로드하여 Asset 생성 (SDK 4.2 권장 방식)
        with open(pdf_path, 'rb') as file:
            input_stream = file.read()
        input_asset = pdf_services.upload(input_stream=input_stream, mime_type='application/pdf')
        print("✅ Asset 생성 성공")
        
        # Extract 작업 생성
        extract_pdf_job = ExtractPDFJob(
            input_asset=input_asset, 
            extract_pdf_params=extract_pdf_params
        )
        
        print("⏳ Adobe Extract 작업 실행 중...")
        
        # 작업 제출 및 결과 대기 (SDK 4.2 호환성 개선)
        try:
            location = pdf_services.submit(extract_pdf_job)
            pdf_services_response = pdf_services.get_job_result(location, ExtractPDFResult)
        except ServiceApiException as api_error:
            print(f"❌ Adobe API ServiceApiException 발생:")
            print(f"   - 에러 메시지: {api_error}")
            print(f"   - 에러 타입: {type(api_error)}")
            if hasattr(api_error, 'status_code'):
                print(f"   - HTTP 상태 코드: {api_error.status_code}")
            if hasattr(api_error, 'error_code'):
                print(f"   - Adobe 에러 코드: {api_error.error_code}")
            if hasattr(api_error, 'message'):
                print(f"   - 상세 메시지: {api_error.message}")
            
            # API 관련 오류 처리
            if "No result class found" in str(api_error) or "no extractable content" in str(api_error).lower():
                print("💡 추출 가능한 텍스트가 없는 PDF (스캔된 이미지) - OCR 백업 모드로 전환")
                return None
            return None
        except ServiceUsageException as usage_error:
            print(f"❌ Adobe 사용량 오류: {usage_error}")
            return None
        except Exception as submit_error:
            print(f"❌ Adobe 작업 제출 오류: {submit_error}")
            # 파일이 스캔된 이미지인 경우 OCR 모드로 재시도
            if "No result class found" in str(submit_error) or "invalid" in str(submit_error).lower():
                print("💡 스캔된 PDF로 판단 - OCR 백업 모드로 전환")
                return None
            return None
        
        # 결과 확인
        if not pdf_services_response:
            print("❌ Adobe Extract 응답이 없습니다")
            return None
            
        result = pdf_services_response.get_result()
        if not result:
            print("❌ Adobe Extract 결과가 없습니다")
            return None
            
        result_asset = result.get_resource()
        if not result_asset:
            print("❌ Adobe Extract 결과 에셋이 없습니다")
            return None
        
        # 결과를 임시 파일로 저장
        temp_dir = tempfile.mkdtemp()
        temp_zip_path = os.path.join(temp_dir, "extract_result.zip")
        
        # 결과 스트림을 파일로 저장 (SDK 4.2 올바른 방식)
        download_success = False
        
        # SDK 4.2의 올바른 방법: StreamAsset의 get_input_stream() 사용
        try:
            if hasattr(result_asset, 'get_input_stream'):
                # StreamAsset의 경우 - bytes-like object 처리
                input_stream = result_asset.get_input_stream()
                
                # 스트림 데이터를 bytes로 변환
                if hasattr(input_stream, 'read'):
                    # 파일 객체인 경우
                    stream_data = input_stream.read()
                elif hasattr(input_stream, '__enter__'):
                    # 컨텍스트 매니저인 경우
                    with input_stream as stream:
                        stream_data = stream.read()
                else:
                    # 이미 bytes인 경우
                    stream_data = input_stream
                
                # bytes 타입 확인 및 변환
                if isinstance(stream_data, bytes):
                    with open(temp_zip_path, 'wb') as temp_file:
                        temp_file.write(stream_data)
                    download_success = True
                    print("✅ StreamAsset.get_input_stream() 성공")
                else:
                    # str이나 다른 타입인 경우 bytes로 변환
                    if isinstance(stream_data, str):
                        stream_data = stream_data.encode('utf-8')
                    elif hasattr(stream_data, 'encode'):
                        stream_data = stream_data.encode()
                    
                    with open(temp_zip_path, 'wb') as temp_file:
                        temp_file.write(stream_data)
                    download_success = True
                    print("✅ StreamAsset 데이터 변환 후 저장 성공")
            else:
                raise AttributeError("result_asset에 get_input_stream 메서드가 없습니다")
        except Exception as stream_error:
            print(f"⚠️ StreamAsset 다운로드 실패: {stream_error}")
        
        # 대안 방법: CloudAsset 및 기타 Asset 타입 처리
        if not download_success:
            try:
                # CloudAsset이나 다른 타입의 경우
                if hasattr(result_asset, 'download_uri'):
                    # CloudAsset의 download_uri 사용
                    import requests
                    response = requests.get(result_asset.download_uri)
                    response.raise_for_status()
                    
                    # bytes 데이터 확인 및 저장
                    content_data = response.content
                    if isinstance(content_data, bytes):
                        with open(temp_zip_path, 'wb') as temp_file:
                            temp_file.write(content_data)
                        download_success = True
                        print("✅ CloudAsset download_uri 성공")
                    else:
                        print(f"⚠️ CloudAsset 응답이 bytes가 아님: {type(content_data)}")
                        
                elif hasattr(result_asset, 'get_stream'):
                    # 다른 스트림 메서드 시도 - bytes 처리 강화
                    stream_data = result_asset.get_stream()
                    
                    # 스트림 데이터 타입별 처리
                    if hasattr(stream_data, 'read'):
                        # 파일 객체인 경우
                        raw_data = stream_data.read()
                    else:
                        # 직접 데이터인 경우
                        raw_data = stream_data
                    
                    # bytes 변환 및 저장
                    if isinstance(raw_data, bytes):
                        final_data = raw_data
                    elif isinstance(raw_data, str):
                        final_data = raw_data.encode('utf-8')
                    elif hasattr(raw_data, 'encode'):
                        final_data = raw_data.encode()
                    else:
                        # 다른 타입인 경우 str로 변환 후 bytes로
                        final_data = str(raw_data).encode('utf-8')
                    
                    with open(temp_zip_path, 'wb') as temp_file:
                        temp_file.write(final_data)
                    download_success = True
                    print("✅ Asset get_stream 성공")
                    
                else:
                    raise AttributeError("Asset에서 데이터를 가져올 수 있는 메서드를 찾을 수 없습니다")
            except Exception as alt_error:
                print(f"⚠️ 대안 다운로드 방법 실패: {alt_error}")
        
        # 모든 방법 실패 시 오류 반환
        if not download_success:
            print("❌ 모든 결과 다운로드 방법 실패")
            print(f"📋 result_asset 타입: {type(result_asset)}")
            print(f"📋 result_asset 속성: {dir(result_asset)}")
            return None
        
        print(f"📥 Adobe 결과 다운로드 완료: {os.path.getsize(temp_zip_path) / 1024:.1f}KB")
        
        # ZIP에서 structuredData.json 추출 및 파싱
        page_blocks = []
        
        try:
            with zipfile.ZipFile(temp_zip_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                print(f"📋 Adobe 결과 파일 목록: {file_list}")
                
                if 'structuredData.json' in file_list:
                    with zip_ref.open('structuredData.json') as json_file:
                        data = json.load(json_file)
                        page_blocks = parse_adobe_elements(data)
                        print(f"✅ Adobe 텍스트 추출 성공: {len(page_blocks)} 페이지")
                else:
                    print("⚠️ structuredData.json이 Adobe 응답에 없습니다")
                    # 다른 JSON 파일 확인
                    json_files = [f for f in file_list if f.endswith('.json')]
                    if json_files:
                        print(f"📄 발견된 JSON 파일들: {json_files}")
        except zipfile.BadZipFile:
            print("❌ Adobe 응답이 유효한 ZIP 파일이 아닙니다")
        
        # 임시 파일 정리 (안전한 정리)
        try:
            if os.path.exists(temp_zip_path):
                os.unlink(temp_zip_path)
            if os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as cleanup_error:
            print(f"⚠️ 임시 파일 정리 오류: {cleanup_error}")
        
        if page_blocks:
            print(f"✅ Adobe Extract 완료: {len(page_blocks)} 페이지, 총 {sum(len(blocks) for blocks in page_blocks)} 텍스트 블록")
            return page_blocks
        else:
            print("⚠️ Adobe Extract에서 텍스트를 추출하지 못했습니다")
            return None
        
    except Exception as e:
        error_msg = str(e)
        print(f"❌ Adobe Extract 오류: {error_msg}")
        
        # 구체적인 오류 메시지 분석 및 복구 가이드
        if "No result class found" in error_msg:
            print("💡 'No result class found' 오류 - PDF 파일이 텍스트를 포함하지 않거나 스캔된 이미지일 수 있습니다")
        elif "credentials" in error_msg.lower() or "authentication" in error_msg.lower():
            print("💡 인증 오류 - Adobe API 키를 확인해주세요")
        elif "timeout" in error_msg.lower() or "connection" in error_msg.lower():
            print("💡 네트워크 오류 - 인터넷 연결 또는 Adobe 서버 상태를 확인해주세요")
        elif "file" in error_msg.lower() and "corrupt" in error_msg.lower():
            print("💡 파일 손상 오류 - PDF 파일이 손상되었을 수 있습니다")
        elif "memory" in error_msg.lower() or "out of memory" in error_msg.lower():
            print("💡 메모리 부족 오류 - 파일 크기를 줄이거나 시스템 메모리를 확인해주세요")
        
        # 임시 파일 정리 (오류 발생 시에도)
        try:
            if 'temp_zip_path' in locals() and os.path.exists(temp_zip_path):
                os.unlink(temp_zip_path)
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
        
        return None

def parse_adobe_elements(data):
    """Adobe structuredData.json을 페이지별 텍스트 블록으로 파싱 (하이브리드 모드: 텍스트 + 벡터 이미지)"""
    page_blocks = []
    
    try:
        elements = data.get('elements', [])
        current_page_blocks = []
        current_page = 0
        
        for element in elements:
            if element.get('Page') != current_page:
                if current_page_blocks:
                    page_blocks.append(current_page_blocks)
                current_page_blocks = []
                current_page = element.get('Page', 0)
            
            # 텍스트 요소 처리
            if element.get('Path') and element.get('Text'):  
                bounds = element.get('Bounds', [])
                text = element.get('Text', '').strip()
                
                if text and len(bounds) >= 4:
                    block = {
                        'type': 'text',
                        'text': text,
                        'left': bounds[0],
                        'top': bounds[1], 
                        'width': bounds[2] - bounds[0],
                        'height': bounds[3] - bounds[1],
                        'confidence': 100  # Adobe는 높은 신뢰도로 가정
                    }
                    current_page_blocks.append(block)
            
            # 이미지/그림 요소 처리 (벡터로 변환)
            elif element.get('Path') and '/Figure' in element.get('Path', ''):
                bounds = element.get('Bounds', [])
                file_paths = element.get('filePaths', [])
                
                if len(bounds) >= 4:
                    # 벡터 이미지 블록 생성
                    vector_block = {
                        'type': 'vector_image',
                        'text': '[벡터 이미지]',  # 편집 가능한 텍스트 표현
                        'left': bounds[0],
                        'top': bounds[1],
                        'width': bounds[2] - bounds[0],
                        'height': bounds[3] - bounds[1],
                        'confidence': 100,
                        'file_paths': file_paths,
                        'vector_description': f'이미지 영역 ({bounds[2] - bounds[0]:.0f}x{bounds[3] - bounds[1]:.0f}px)'
                    }
                    current_page_blocks.append(vector_block)
                    print(f"🎨 벡터 이미지 요소 추가: {bounds[2] - bounds[0]:.0f}x{bounds[3] - bounds[1]:.0f}px")
            
            # 테이블 요소 처리
            elif element.get('Path') and '/Table' in element.get('Path', ''):
                bounds = element.get('Bounds', [])
                
                if len(bounds) >= 4:
                    table_block = {
                        'type': 'table',
                        'text': '[표 영역]',  # 편집 가능한 텍스트 표현
                        'left': bounds[0],
                        'top': bounds[1],
                        'width': bounds[2] - bounds[0],
                        'height': bounds[3] - bounds[1],
                        'confidence': 100,
                        'table_description': f'표 영역 ({bounds[2] - bounds[0]:.0f}x{bounds[3] - bounds[1]:.0f}px)'
                    }
                    current_page_blocks.append(table_block)
                    print(f"📊 테이블 요소 추가: {bounds[2] - bounds[0]:.0f}x{bounds[3] - bounds[1]:.0f}px")
        
        # 마지막 페이지 추가
        if current_page_blocks:
            page_blocks.append(current_page_blocks)
            
        # 하이브리드 처리 결과 요약
        total_text = sum(1 for page in page_blocks for block in page if block.get('type') == 'text')
        total_vectors = sum(1 for page in page_blocks for block in page if block.get('type') == 'vector_image')
        total_tables = sum(1 for page in page_blocks for block in page if block.get('type') == 'table')
        print(f"🎯 하이브리드 파싱 완료: 텍스트 {total_text}개, 벡터 이미지 {total_vectors}개, 테이블 {total_tables}개")
            
    except Exception as e:
        print(f"Adobe 데이터 파싱 오류: {e}")
        
    return page_blocks

def add_editable_text_with_adobe(doc, image, section, adobe_blocks):
    """Adobe Extract로 추출한 텍스트를 완전히 편집 가능한 순수 텍스트로 Word 문서에 추가 (하이브리드 모드: 텍스트 + 벡터 이미지)"""
    try:
        print(f"  - 📝 Adobe 하이브리드 편집 가능 변환 시작: {len(adobe_blocks)}개 블록")
        
        if not adobe_blocks:
            print("  - ⚠️ Adobe 텍스트 블록이 없어 빈 문서 생성")
            para = doc.add_paragraph("Adobe SDK에서 텍스트를 추출할 수 없었습니다.")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        
        # 블록 타입별 통계
        text_blocks = [b for b in adobe_blocks if b.get('type') == 'text']
        vector_blocks = [b for b in adobe_blocks if b.get('type') == 'vector_image']
        table_blocks = [b for b in adobe_blocks if b.get('type') == 'table']
        
        print(f"  - 🎯 하이브리드 처리: 텍스트 {len(text_blocks)}개, 벡터 이미지 {len(vector_blocks)}개, 테이블 {len(table_blocks)}개")
        
        # 텍스트 블록을 Y 좌표 기준으로 정렬 (위에서 아래로)
        sorted_blocks = sorted(adobe_blocks, key=lambda x: (x.get('top', 0), x.get('left', 0)))
        
        # 이미지 크기 정보 (좌표 변환용)
        img_width, img_height = image.size if image else (1, 1)
        
        print(f"  - 📄 순수 텍스트 모드: 이미지 배경 없이 완전 편집 가능한 텍스트만 생성")
        print(f"  - 📊 원본 이미지 크기: {img_width}x{img_height}px (좌표 변환 참조용)")
        
        # 문서 제목 추가 (첫 번째 텍스트 블록이 제목인 경우)
        if sorted_blocks:
            first_block = sorted_blocks[0]
            first_text = first_block.get('text', '').strip()
            
            # 제목으로 보이는 첫 번째 텍스트 처리
            if any(keyword in first_text for keyword in ['공문', '통지', '안내', '요청', '회신', '발명']):
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(first_text)
                title_run.font.name = "맑은 고딕"
                title_run.font.size = Pt(14)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0, 0, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.paragraph_format.space_after = Pt(12)
                
                # 제목 처리했으므로 나머지 블록만 처리
                sorted_blocks = sorted_blocks[1:]
                print(f"  - 📋 문서 제목 설정: '{first_text[:30]}...'")
        
        # 하이브리드 블록을 순수 편집 가능한 텍스트로 변환 (텍스트 + 벡터 이미지 + 테이블)
        for i, block in enumerate(sorted_blocks):
            text_content = block.get('text', '').strip()
            block_type = block.get('type', 'text')
            
            if not text_content:
                continue
            
            # 새 문단 생성 (완전 편집 가능)
            para = doc.add_paragraph()
            run = para.add_run(text_content)
            
            # 블록 타입별 서식 설정
            if block_type == 'text':
                # 일반 텍스트 블록
                run.font.name = "맑은 고딕"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
            elif block_type == 'vector_image':
                # 벡터 이미지 블록 - 편집 가능한 텍스트로 표현
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(51, 102, 204)  # 파란색
                run.font.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 벡터 설명 추가
                if block.get('vector_description'):
                    desc_run = para.add_run(f"\n{block['vector_description']}")
                    desc_run.font.name = "맑은 고딕"
                    desc_run.font.size = Pt(9)
                    desc_run.font.color.rgb = RGBColor(102, 102, 102)  # 회색
                    desc_run.font.italic = True
                
            elif block_type == 'table':
                # 테이블 블록 - 편집 가능한 텍스트로 표현
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(204, 102, 51)  # 주황색
                run.font.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 테이블 설명 추가
                if block.get('table_description'):
                    desc_run = para.add_run(f"\n{block['table_description']}")
                    desc_run.font.name = "맑은 고딕"
                    desc_run.font.size = Pt(9)
                    desc_run.font.color.rgb = RGBColor(102, 102, 102)  # 회색
                    desc_run.font.italic = True
            
            # Adobe 좌표를 기반으로 한 레이아웃 추정
            left_ratio = block.get('left', 0) / img_width if img_width > 0 else 0
            
            # 텍스트 정렬 및 서식 결정 (좌표 기반)
            if left_ratio > 0.75:  # 우측 정렬 (날짜, 서명 등)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run.font.size = Pt(10)
            elif 0.25 <= left_ratio <= 0.75:  # 중앙 정렬 (제목, 부제목 등)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if any(keyword in text_content for keyword in ['제목', 'MARS', 'CONTEST']):
                    run.font.bold = True
                    run.font.size = Pt(12)
            else:  # 좌측 정렬 (본문)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 들여쓰기 적용 (본문 구조 반영)
                if text_content.startswith(('1.', '2.', '3.', '가.', '나.', '다.')):
                    para.paragraph_format.left_indent = Pt(18)
                elif text_content.startswith(('•', '-', '○')):
                    para.paragraph_format.left_indent = Pt(36)
            
            # 특수 텍스트 서식 적용
            if any(keyword in text_content for keyword in ['수신', '발신', '제목', '내용']):
                run.font.bold = True
            elif any(keyword in text_content for keyword in ['연구소', '기관', '부서']):
                run.font.size = Pt(12)
                run.font.bold = True
            elif text_content.startswith('붙임'):
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 문단 간격 설정 (자연스러운 레이아웃)
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(6)
            
            # Y 좌표 기반 상단 여백 조정
            if i > 0:
                prev_block = sorted_blocks[i-1]
                prev_bottom = prev_block.get('top', 0) + prev_block.get('height', 0)
                current_top = block.get('top', 0)
                y_gap = current_top - prev_bottom
                
                # 큰 간격이 있는 경우 문단 간격 추가
                if y_gap > 30:
                    para.paragraph_format.space_before = Pt(8)
                elif y_gap > 15:
                    para.paragraph_format.space_before = Pt(4)
            
            # 완전 편집 가능하도록 표준 스타일 적용
            para.style = doc.styles['Normal']
        
        # 문서 하단에 편집 안내 추가
        footer_para = doc.add_paragraph()
        footer_para.paragraph_format.space_before = Pt(20)
        footer_run = footer_para.add_run("※ 이 문서의 모든 텍스트는 완전히 편집 가능합니다.")
        footer_run.font.name = "맑은 고딕"
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        print(f"  - ✅ Adobe 순수 텍스트 변환 완료: {len(sorted_blocks)}개 블록을 완전 편집 가능한 텍스트로 변환")
        print(f"  - 🎯 편집성: 100% (이미지 배경 없음, 순수 텍스트만)")
        return True
        
    except Exception as e:
        print(f"  - ❌ Adobe 순수 텍스트 변환 오류: {e}")
        return False

def add_image_with_adobe_text(doc, image, section, adobe_blocks):
    """하위 호환성을 위한 래퍼 함수 - 새로운 편집 가능 텍스트 변환 사용"""
    return add_editable_text_with_adobe(doc, image, section, adobe_blocks)

def setup_tesseract():
    """Tesseract OCR 설정"""
    try:
        tesseract_paths = [
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"Tesseract 경로 설정 성공: {path}")
                return True
        # PATH 탐지
        try:
            pytesseract.get_tesseract_version()
            print("Tesseract가 시스템 PATH에서 발견됨")
            return True
        except Exception:
            print("⚠️ Tesseract OCR을 찾을 수 없습니다."
                  " 텍스트 편집 기능은 제한됩니다.")
            return False
    except Exception as e:
        print(f"Tesseract 설정 오류: {e}")
        return False

def ocr_image_to_blocks(pil_image):
    """이미지에서 단어 단위 텍스트와 위치(좌표)를 추출"""
    try:
        img = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.medianBlur(gray, 3)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray = clahe.apply(gray)

        config = r"--oem 3 --psm 6 -l kor+eng"
        data = pytesseract.image_to_data(gray, config=config,
                                         output_type=pytesseract.Output.DICT)
        blocks = []
        n = len(data["text"])
        for i in range(n):
            text = data["text"][i].strip()
            # conf 값 안전하게 처리 (정수, 문자열, 실수 모두 고려)
            conf_val = data["conf"][i]
            if isinstance(conf_val, (int, float)):
                conf = int(conf_val)
            elif isinstance(conf_val, str) and conf_val.replace('.', '').replace('-', '').isdigit():
                conf = int(float(conf_val))
            else:
                conf = -1
            if text and conf >= 0:
                blocks.append({
                    "text": text,
                    "x": data["left"][i],
                    "y": data["top"][i],
                    "w": data["width"][i],
                    "h": data["height"][i],
                    "conf": conf,
                })
        return blocks
    except Exception as e:
        error_msg = str(e)
        print(f"❌ OCR 블록 추출 오류: {error_msg}")
        
        # 구체적인 오류 분석 및 복구 가이드
        if "tesseract" in error_msg.lower():
            print("💡 Tesseract 오류 - OCR 엔진이 설치되지 않았거나 언어팩이 누락되었을 수 있습니다")
        elif "memory" in error_msg.lower() or "out of memory" in error_msg.lower():
            print("💡 메모리 부족 - 이미지 크기를 줄이거나 시스템 메모리를 확인해주세요")
        elif "timeout" in error_msg.lower():
            print("💡 처리 시간 초과 - 이미지가 너무 복잡하거나 큽니다")
        elif "opencv" in error_msg.lower() or "cv2" in error_msg.lower():
            print("💡 OpenCV 오류 - 이미지 전처리 중 문제가 발생했습니다")
        
        # 메모리 정리
        try:
            if 'processed_image' in locals():
                del processed_image
            if 'img_array' in locals():
                del img_array
            import gc
            gc.collect()
        except Exception:
            pass
        
        return []

def clean_korean_text(text):
    """한글 공문서 특화 텍스트 정제 함수"""
    try:
        import re
        
        if not text or not text.strip():
            return ""
        
        # 1. 기본 정제
        cleaned = text.strip()
        
        # 2. OCR 오인식 패턴 수정 (한글 공문서 특화)
        # 자주 오인식되는 한글 문자 패턴 수정
        ocr_corrections = {
            'ㅇ': '○',  # 원 기호 오인식
            'ㅁ': '□',  # 사각형 기호 오인식
            'l': '1',   # 소문자 l과 숫자 1 구분
            'O': '0',   # 대문자 O와 숫자 0 구분 (맥락에 따라)
            '|': '1',   # 세로선과 숫자 1 구분
        }
        
        # 3. 연속된 공백 정리
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        # 4. 특수문자 정리 (한글 공문서에서 의미있는 문자만 보존)
        # 불필요한 특수문자 제거 (단, 공문서에서 사용되는 기호는 보존)
        cleaned = re.sub(r'[^가-힣ㄱ-ㅎㅏ-ㅣa-zA-Z0-9\s()\[\]{}.,?!\-+=:;"\'\/·※○●△▲▼◆■□◇◎★☆]', '', cleaned)
        
        # 5. 최종 검증
        if len(cleaned.strip()) == 0:
            return ""
        
        # 6. 의미있는 문자 비율 검사 (한글/영문/숫자가 50% 이상)
        meaningful_chars = len([c for c in cleaned if c.isalnum() or c in '가-힣ㄱ-ㅎㅏ-ㅣ'])
        total_chars = len(cleaned.replace(' ', ''))
        
        if total_chars > 0 and meaningful_chars / total_chars >= 0.5:
            return cleaned.strip()
        else:
            return ""
            
    except Exception as e:
        print(f"텍스트 정제 오류: {e}")
        return text.strip() if text else ""

def extract_text_blocks_with_ocr(image):
    """OCR을 사용하여 이미지에서 텍스트 블록 추출 (개선된 버전)"""
    try:
        # 이미지 전처리로 OCR 정확도 향상
        import cv2
        import numpy as np
        
        # PIL 이미지를 OpenCV 형식으로 변환
        img_array = np.array(image)
        
        # 이미지 전처리
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        
        # 한글 공문서 최적화된 이미지 전처리 (정확도 향상)
        # 1. 이미지 크기 정규화 (OCR 최적 해상도로 조정)
        height, width = gray.shape
        if height < 300 or width < 300:  # 너무 작은 이미지는 확대
            scale_factor = max(300 / height, 300 / width)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            gray = cv2.resize(gray, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
        
        # 2. 노이즈 제거 (한글 문자 보존 강화)
        denoised = cv2.fastNlMeansDenoising(gray, h=8, templateWindowSize=7, searchWindowSize=21)
        
        # 3. 대비 향상 (CLAHE 적용 - 한글 최적화)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # 4. 가우시안 블러로 미세한 노이즈 제거
        # OpenCV Gaussian kernel 오류 방지: ksize 조건 검증
        blur_ksize = (3, 3)  # (1,1)은 너무 작으므로 (3,3)으로 변경
        if blur_ksize[0] > 0 and blur_ksize[0] % 2 == 1 and blur_ksize[1] > 0 and blur_ksize[1] % 2 == 1:
            blurred = cv2.GaussianBlur(enhanced, blur_ksize, 0)
        else:
            blurred = enhanced.copy()  # 블러 없이 원본 사용
        
        # 5. 언샤프 마스킹으로 텍스트 선명도 향상
        unsharp_strength = 1.5
        # (0,0) ksize는 sigmaX, sigmaY로 자동 계산되므로 유효함
        gaussian = cv2.GaussianBlur(blurred, (0, 0), 2.0)
        sharpened = cv2.addWeighted(blurred, 1.0 + unsharp_strength, gaussian, -unsharp_strength, 0)
        
        # 6. 적응형 임계값으로 이진화 (한글 최적화 파라미터)
        thresh = cv2.adaptiveThreshold(sharpened, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # 7. 모폴로지 연산으로 문자 연결성 개선 (한글 특성 고려)
        kernel_close = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 1))  # 가로 연결
        closed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel_close)
        
        # 8. 작은 노이즈 제거 (한글 문자는 보존)
        kernel_open = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (2, 2))
        processed = cv2.morphologyEx(closed, cv2.MORPH_OPEN, kernel_open)
        
        # 전처리된 이미지로 OCR 수행
        processed_image = Image.fromarray(processed)
        
        # 한글 공문서 최적화된 Tesseract 설정 (정확도 향상)
        # PSM 3: 완전 자동 페이지 분할 (공문서 레이아웃 최적화)
        # OEM 1: LSTM OCR 엔진만 사용 (한글 인식률 최대화)
        # 한글 공문서 특화 설정
        config = r'--oem 1 --psm 3 -l kor+eng -c tessedit_char_whitelist=가-힣ㄱ-ㅎㅏ-ㅣ0-9A-Za-z()[]{}.,?!-+=:;"\'\'\ /\n\t·※○●△▲▼◆■□◇◎★☆ -c preserve_interword_spaces=1 -c tessedit_do_invert=0'
        
        # OCR 수행 (타임아웃 및 오류 처리 강화)
        try:
            data = pytesseract.image_to_data(processed_image, config=config, output_type=pytesseract.Output.DICT, timeout=30)
        except pytesseract.TesseractError as te:
            print(f"  - ⚠️ Tesseract 설정 오류, 기본 설정으로 재시도: {te}")
            # 기본 설정으로 재시도
            try:
                data = pytesseract.image_to_data(processed_image, lang='kor+eng', output_type=pytesseract.Output.DICT, timeout=30)
            except Exception as retry_error:
                print(f"  - ❌ OCR 재시도 실패: {retry_error}")
                return []
        except Exception as ocr_error:
            print(f"  - ❌ OCR 처리 오류: {ocr_error}")
            return []
        
        blocks = []
        
        # 텍스트 블록을 라인별로 그룹화 (한글 공문서 최적화)
        lines = {}
        valid_texts = []
        
        # 1단계: 유효한 텍스트만 필터링 (신뢰도 및 품질 기준 강화)
        for i in range(len(data['text'])):
            text = data['text'][i].strip()
            conf = int(data['conf'][i])
            
            # 한글 공문서 특화 필터링 조건
            if (conf > 30 and text and  # 신뢰도 30% 이상으로 상향
                len(text) >= 1 and  # 1글자 이상 (한글 특성 고려)
                not text.isspace() and  # 공백만 있는 텍스트 제외
                len([c for c in text if c.isalnum() or c in '가-힣ㄱ-ㅎㅏ-ㅣ']) > 0):  # 의미있는 문자 포함
                
                valid_texts.append({
                    'text': text,
                    'left': data['left'][i],
                    'top': data['top'][i],
                    'width': data['width'][i],
                    'height': data['height'][i],
                    'confidence': conf
                })
        
        # 2단계: 라인별 그룹화 (한글 공문서 레이아웃 고려)
        for item in valid_texts:
            text = item['text']
            left = item['left']
            top = item['top']
            width = item['width']
            height = item['height']
            conf = item['confidence']
            
            # 동적 라인 그룹화 (텍스트 높이 기준)
            line_tolerance = max(8, height // 3)  # 텍스트 높이의 1/3 또는 최소 8픽셀
            line_key = round(top / line_tolerance) * line_tolerance
            
            if line_key not in lines:
                lines[line_key] = {
                    'texts': [],
                    'positions': [],
                    'top': top,
                    'left': left,
                    'width': width,
                    'height': height,
                    'confidence': conf
                }
            
            lines[line_key]['texts'].append(text)
            lines[line_key]['positions'].append({'left': left, 'text': text})
            lines[line_key]['left'] = min(lines[line_key]['left'], left)
            lines[line_key]['width'] = max(lines[line_key]['width'], left + width - lines[line_key]['left'])
            lines[line_key]['height'] = max(lines[line_key]['height'], height)
            lines[line_key]['confidence'] = max(lines[line_key]['confidence'], conf)
        
        # 3단계: 라인별 블록 생성 (한글 공문서 텍스트 순서 보존)
        for line_key, line_data in sorted(lines.items()):
            if line_data['texts'] and line_data['positions']:
                # 같은 라인 내에서 좌측부터 정렬 (한글 공문서 읽기 순서)
                sorted_positions = sorted(line_data['positions'], key=lambda x: x['left'])
                ordered_texts = [pos['text'] for pos in sorted_positions]
                
                # 텍스트 결합 (한글 공문서 특성 고려)
                combined_text = ' '.join(ordered_texts).strip()
                
                # 품질 검증 및 블록 생성
                if (len(combined_text) >= 1 and  # 최소 1글자 이상
                    not combined_text.isspace() and  # 공백만 있는 텍스트 제외
                    line_data['confidence'] > 25):  # 신뢰도 25% 이상
                    
                    # 한글 공문서 특화 텍스트 정제
                    cleaned_text = clean_korean_text(combined_text)
                    
                    if cleaned_text:  # 정제 후에도 유효한 텍스트가 있는 경우
                        blocks.append({
                            'left': line_data['left'],
                            'top': line_data['top'],
                            'width': line_data['width'],
                            'height': line_data['height'],
                            'confidence': line_data['confidence'],
                            'text': cleaned_text
                        })
        
        print(f"  - OCR 텍스트 블록 {len(blocks)}개 추출됨")
        for i, block in enumerate(blocks[:3]):  # 처음 3개만 로그 출력
            print(f"    블록 {i+1}: '{block['text'][:30]}...' (신뢰도: {block['confidence']}%)")
        
        return blocks
        
    except Exception as e:
        print(f"OCR 블록 추출 오류: {e}")
        return []

def add_image_and_overlay_text(doc, image, section):
    """스마트 문서 타입 감지 및 적응형 변환"""
    try:
        print("  - 🔍 문서 타입 감지 시작")
        
        # OCR 텍스트 추출
        text_blocks = extract_text_blocks_with_ocr(image)
        print(f"  - OCR 텍스트 블록 {len(text_blocks)}개 감지")
        
        # 문서 타입 감지
        doc_type = detect_document_type(image, text_blocks)
        
        # 문서 타입에 따른 변환 모드 선택
        if doc_type == "text_only":
            # 2번 이미지: 텍스트만 편집 가능
            success = add_text_only_conversion(doc, image, section, text_blocks)
        elif doc_type == "hybrid":
            # 3번 이미지: 이미지+텍스트 블럭 편집 가능
            success = add_hybrid_conversion(doc, image, section, text_blocks)
        else:  # image_only
            # 4번 이미지: 이미지 블럭 그대로 유지
            success = add_image_only_conversion(doc, image, section)
        
        if success:
            print(f"  - ✅ {doc_type} 모드 변환 완료")
        else:
            print(f"  - ❌ {doc_type} 모드 변환 실패")
            
        return success
        
    except Exception as e:
        print(f"  - ❌ 스마트 변환 오류: {e}")
        # 오류 시 기본 이미지만 추가
        return add_image_only_conversion(doc, image, section)

def detect_image_orientation(image):
    """이미지 방향 감지 (가로/세로) - A4 표준 비율 기준"""
    try:
        width, height = image.size
        aspect_ratio = width / height
        
        # A4 표준 비율 (297/210 ≈ 1.414)
        a4_landscape_ratio = 297 / 210  # 가로형 A4 비율
        a4_portrait_ratio = 210 / 297   # 세로형 A4 비율
        
        print(f"  - 이미지 크기: {width} x {height} 픽셀")
        print(f"  - 이미지 비율: {aspect_ratio:.3f}")
        print(f"  - A4 가로형 비율: {a4_landscape_ratio:.3f}")
        print(f"  - A4 세로형 비율: {a4_portrait_ratio:.3f}")
        
        # A4 표준 비율과 비교하여 방향 판단
        if aspect_ratio >= 1.3:  # A4 가로형에 가까운 비율
            print(f"  - ✅ 가로형 감지됨 (비율: {aspect_ratio:.3f} >= 1.3)")
            return "landscape"
        elif aspect_ratio <= 0.8:  # A4 세로형에 가까운 비율
            print(f"  - ✅ 세로형 감지됨 (비율: {aspect_ratio:.3f} <= 0.8)")
            return "portrait"
        else:  # 중간 비율인 경우 - A4 표준과 더 가까운 쪽으로 판단
            landscape_diff = abs(aspect_ratio - a4_landscape_ratio)
            portrait_diff = abs(aspect_ratio - a4_portrait_ratio)
            
            if landscape_diff < portrait_diff:
                print(f"  - ✅ A4 가로형에 더 가까움, 가로형으로 처리 (비율: {aspect_ratio:.3f})")
                return "landscape"
            else:
                print(f"  - ✅ A4 세로형에 더 가까움, 세로형으로 처리 (비율: {aspect_ratio:.3f})")
                return "portrait"
    except Exception as e:
        print(f"이미지 방향 감지 오류: {e}")
        return "landscape"  # 기본값을 가로로 변경

def _set_section_orientation(section, orientation: str):
    """섹션 용지 방향과 크기를 A4에 맞춰 설정한다."""
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    margin = Mm(15)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

def _fit_dimensions_within(max_width_inch: float, max_height_inch: float, img_width_inch: float, img_height_inch: float):
    """주어진 영역 안에 비율을 유지하며 이미지를 맞춘다."""
    scale = min(max_width_inch / img_width_inch, max_height_inch / img_height_inch)
    return Inches(img_width_inch * scale), Inches(img_height_inch * scale)

def pdf_to_docx(pdf_path, output_path):
    """PDF를 DOCX로 변환 - Adobe PDF Services SDK 4.2 무조건 우선 사용.
    Adobe SDK를 통한 완전 편집 가능한 텍스트 추출을 최우선으로 처리하고,
    SDK 실패 시에만 OCR 백업 사용. (안정성 강화)"""
    images = None
    doc = None
    try:
        print(f"PDF → DOCX 변환 시작: {pdf_path}")
        
        # 파일 존재 및 크기 확인
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {pdf_path}")
        
        file_size = os.path.getsize(pdf_path) / (1024 * 1024)  # MB
        print(f"📄 PDF 파일 크기: {file_size:.1f}MB")
        
        if file_size > 100:  # 100MB 이상
            print(f"⚠️ 큰 파일 감지 ({file_size:.1f}MB) - 처리 시간이 오래 걸릴 수 있습니다")
        
        # 1) Adobe PDF Services SDK ExportPDFOperation 우선 사용 (직접 DOCX 변환) - 1회만 시도
        adobe_success = False
        
        if ADOBE_SDK_AVAILABLE:
            try:
                print(f"🔗 Adobe PDF Services SDK ExportPDFOperation 우선 사용 시작...")
                
                # Adobe ExportPDF API로 직접 DOCX 변환 (1회만 시도)
                adobe_success = convert_pdf_to_docx_with_adobe(pdf_path, output_path)
                
                if adobe_success:
                    print(f"✅ Adobe SDK ExportPDF 성공: PDF를 편집 가능한 DOCX로 직접 변환 완료")
                    print(f"📄 변환 완료: {output_path}")
                    return True
                else:
                    print(f"⚠️ Adobe SDK ExportPDF 실패 - Extract API 백업으로 전환")
                    
            except Exception as e:
                print(f"❌ Adobe SDK ExportPDF 오류: {e}")
                print("🔄 Extract API 백업으로 전환...")
                    
            if not adobe_success:
                print(f"❌ Adobe SDK ExportPDF 실패 - OCR 백업으로 전환")
        else:
            print("⚠️ Adobe SDK 사용 불가 - OCR 백업 사용")
            
        # Adobe ExportPDF 실패 시에만 기존 Extract 방식으로 백업 처리
        adobe_blocks_per_page = None
        if not adobe_success and ADOBE_SDK_AVAILABLE:
            try:
                print("🔄 Adobe Extract API 백업 시도...")
                adobe_blocks_per_page = extract_with_adobe(pdf_path)
                if adobe_blocks_per_page and len(adobe_blocks_per_page) > 0:
                    total_blocks = sum(len(page_blocks) for page_blocks in adobe_blocks_per_page)
                    text_blocks = sum(1 for page_blocks in adobe_blocks_per_page 
                                    for block in page_blocks if block.get('type') == 'text' and block.get('text', '').strip())
                    if text_blocks > 0:
                        print(f"✅ Adobe Extract 백업 성공: {len(adobe_blocks_per_page)}페이지, {total_blocks}개 블록 추출")
                    else:
                        adobe_blocks_per_page = None
            except Exception as e:
                print(f"❌ Adobe Extract 백업도 실패: {e}")
                adobe_blocks_per_page = None

        # 2) Adobe Extract 백업 성공 시 하이브리드 모드, 실패 시 OCR 백업 모드
        if adobe_blocks_per_page:
            print("🎨 Adobe Extract 하이브리드 모드: 텍스트 + 벡터 이미지 + 테이블 통합 처리")
            print("   - 텍스트: 완전 편집 가능한 순수 텍스트")
            print("   - 벡터 이미지: 고품질 설명 + 편집 가능 텍스트")
            print("   - 테이블: 구조화된 편집 가능 텍스트")
            
            # 새 Word 문서 생성 (하이브리드 모드)
            doc = Document()
            
            # 한글 폰트 설정
            setup_korean_font(doc)
            
            for i, page_blocks in enumerate(adobe_blocks_per_page):
                print(f"페이지 {i+1}/{len(adobe_blocks_per_page)} Adobe 하이브리드 처리 중...")
                
                # 페이지별 블록 타입 통계
                text_count = sum(1 for block in page_blocks if block.get('type') == 'text')
                vector_count = sum(1 for block in page_blocks if block.get('type') == 'vector_image')
                table_count = sum(1 for block in page_blocks if block.get('type') == 'table')
                print(f"  - 블록 구성: 텍스트 {text_count}개, 벡터 {vector_count}개, 테이블 {table_count}개")
                
                # 섹션 설정 (A4 세로형 기본)
                if i == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_section_orientation(section, "portrait")
                
                # Adobe 하이브리드 블록을 편집 가능한 텍스트로 추가
                # 페이지별로 처리하되 첫 번째 페이지의 이미지와 섹션 정보 사용
                if adobe_blocks_per_page:
                    # 첫 번째 페이지 이미지 생성 (좌표 변환용)
                    from pdf2image import convert_from_path
                    temp_images = convert_from_path(pdf_path, dpi=200, first_page=i+1, last_page=i+1)
                    page_image = temp_images[0] if temp_images else None
                    
                    add_editable_text_with_adobe(doc, page_image, section, page_blocks)
                    print(f"  - ✅ Adobe 하이브리드 {len(page_blocks)}개 블록 편집 가능하게 추가")
            
            # Adobe Extract 백업 성공 시 바로 저장하고 반환
            doc.save(output_path)
            print(f"✅ Adobe Extract 하이브리드 변환 완료: {output_path}")
            return True
        else:
            print("🖼️ 하이브리드 모드: 배경 이미지 + OCR 텍스트 오버레이")
            # 이미지 렌더링 (하이브리드 모드용)
            images = convert_from_path(pdf_path, dpi=200)
            
            # 새 Word 문서 생성
            doc = Document()
            
            for i, image in enumerate(images):
                print(f"페이지 {i+1}/{len(images)} 하이브리드 처리 중...")
                
                # 이미지 방향 감지
                orientation = detect_image_orientation(image)
                print(f"  - 이미지 방향: {orientation}")
                
                # 섹션 방향/용지 크기 설정
                if i == 0:
                    section = doc.sections[0]
                else:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_section_orientation(section, orientation)
                
                # 배경 이미지 + OCR 텍스트 오버레이 (편집 가능)
                add_image_and_overlay_text(doc, image, section)
        
        # DOCX 파일 저장
        doc.save(output_path)
        print(f"✅ PDF → DOCX 변환 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ PDF → DOCX 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

def _prevent_text_overlap(text_blocks, image_regions=None, min_distance_pt=15):
    """텍스트 블록 간 겹침 방지 및 이미지 영역과의 충돌 회피 - 개선된 분리 로직"""
    if len(text_blocks) <= 1:
        return text_blocks
    
    # Y 좌표 기준으로 정렬
    sorted_blocks = sorted(text_blocks, key=lambda x: x['top'])
    adjusted_blocks = []
    
    print(f"  - 🔧 텍스트 블록 겹침 방지 처리: {len(sorted_blocks)}개 블록")
    
    for i, block in enumerate(sorted_blocks):
        overlap_detected = False
        image_conflict = False
        
        # 1. 이전 텍스트 블록들과 겹침 확인
        for prev_block in adjusted_blocks:
            # IoU 계산을 위한 겹침 영역 확인
            x_overlap = max(0, min(block['left'] + block['width'], prev_block['left'] + prev_block['width']) - 
                           max(block['left'], prev_block['left']))
            y_overlap = max(0, min(block['top'] + block['height'], prev_block['top'] + prev_block['height']) - 
                           max(block['top'], prev_block['top']))
            
            if x_overlap > 0 and y_overlap > 0:
                # 겹침 영역이 있으면 IoU 계산
                intersection = x_overlap * y_overlap
                block_area = block['width'] * block['height']
                prev_area = prev_block['width'] * prev_block['height']
                union = block_area + prev_area - intersection
                iou = intersection / union if union > 0 else 0
                
                # IoU가 0.15 이상이면 겹침으로 판단 (더 엄격하게)
                if iou > 0.15:
                    overlap_detected = True
                    print(f"    ⚠️ 텍스트 블록 겹침 감지: IoU={iou:.2f}")
                    break
        
        # 2. 이미지 영역과의 충돌 확인 (새로 추가)
        if image_regions and not overlap_detected:
            for img_region in image_regions:
                # 배경 이미지는 제외 (전체 레이아웃)
                if img_region.get('type') == 'background':
                    continue
                    
                # 이미지 영역과 텍스트 블록의 겹침 확인
                img_x_overlap = max(0, min(block['left'] + block['width'], img_region['left'] + img_region['width']) - 
                               max(block['left'], img_region['left']))
                img_y_overlap = max(0, min(block['top'] + block['height'], img_region['top'] + img_region['height']) - 
                               max(block['top'], img_region['top']))
                
                if img_x_overlap > 0 and img_y_overlap > 0:
                    # 이미지 영역과 겹침 비율 계산
                    img_intersection = img_x_overlap * img_y_overlap
                    text_area = block['width'] * block['height']
                    overlap_ratio = img_intersection / text_area if text_area > 0 else 0
                    
                    # 텍스트가 이미지 영역과 30% 이상 겹치면 충돌로 판단
                    if overlap_ratio > 0.3:
                        image_conflict = True
                        print(f"    🖼️ 이미지 영역 충돌 감지: {img_region.get('type', 'unknown')} 영역과 {overlap_ratio:.1%} 겹침")
                        break
        
        # 3. 겹침이나 충돌이 없으면 그대로 추가
        if not overlap_detected and not image_conflict:
            adjusted_blocks.append(block)
        else:
            # 4. 겹침이나 충돌이 있으면 위치 조정
            adjusted_block = block.copy()
            
            if overlap_detected and adjusted_blocks:
                # 텍스트 겹침: 이전 블록 아래로 이동
                prev_block = adjusted_blocks[-1]
                adjusted_block['top'] = prev_block['top'] + prev_block['height'] + min_distance_pt
                print(f"    📝 텍스트 위치 조정: Y={block['top']} → Y={adjusted_block['top']}")
            
            elif image_conflict:
                # 이미지 충돌: 텍스트를 이미지 영역 밖으로 이동
                # 충돌하는 이미지 영역 찾기
                for img_region in image_regions:
                    if img_region.get('type') == 'background':
                        continue
                    
                    img_x_overlap = max(0, min(block['left'] + block['width'], img_region['left'] + img_region['width']) - 
                                   max(block['left'], img_region['left']))
                    img_y_overlap = max(0, min(block['top'] + block['height'], img_region['top'] + img_region['height']) - 
                                   max(block['top'], img_region['top']))
                    
                    if img_x_overlap > 0 and img_y_overlap > 0:
                        # 이미지 영역 아래로 텍스트 이동
                        adjusted_block['top'] = img_region['top'] + img_region['height'] + min_distance_pt
                        print(f"    🔄 이미지 회피 조정: Y={block['top']} → Y={adjusted_block['top']}")
                        break
            
            adjusted_blocks.append(adjusted_block)
    
    print(f"  - ✅ 텍스트 블록 정리 완료: {len(adjusted_blocks)}개 블록 (겹침 해결)")
    return adjusted_blocks

def _calculate_textbox_dimensions(text: str, font_size_pt=12):
    """텍스트 길이에 따른 텍스트박스 크기 계산"""
    # 한글 기준 평균 글자 폭 (pt 단위)
    char_width_pt = font_size_pt * 0.6
    char_height_pt = font_size_pt * 1.2
    
    # 텍스트 길이 계산
    text_length = len(text)
    estimated_width = max(text_length * char_width_pt, 50)  # 최소 50pt
    estimated_height = char_height_pt
    
    return estimated_width, estimated_height

def detect_image_regions(image):
    """이미지에서 실제 이미지 영역만 감지 (텍스트 제외) - 개선된 분리 로직"""
    try:
        import cv2
        import numpy as np
        
        # PIL 이미지를 OpenCV 형식으로 변환
        img_array = np.array(image)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        
        # 실제 이미지 영역 감지 (텍스트가 아닌 그래픽 요소)
        regions = []
        
        height, width = gray.shape
        print(f"  - 🔍 이미지/벡터 영역 분석 시작: {width}x{height}")
        
        # 1. 로고/이미지 영역 감지 (상단 좌측) - 공문서 로고
        logo_height = int(height * 0.35)  # 상단 35% 영역 확대
        logo_width = int(width * 0.6)     # 좌측 60% 영역 확대
        logo_region = gray[0:logo_height, 0:logo_width]
        
        # 로고 영역에서 큰 블록 찾기 (개선된 감지)
        # OpenCV Gaussian kernel 오류 방지: ksize 조건 검증
        ksize = (3, 3)
        if ksize[0] > 0 and ksize[0] % 2 == 1 and ksize[1] > 0 and ksize[1] % 2 == 1:
            logo_blur = cv2.GaussianBlur(logo_region, ksize, 0)
        else:
            logo_blur = logo_region.copy()  # 블러 없이 원본 사용
        logo_edges = cv2.Canny(logo_blur, 15, 60)  # 더 민감한 엣지 감지
        
        # 모폴로지 연산으로 연결된 영역 강화
        kernel = np.ones((4, 4), np.uint8)
        logo_edges = cv2.morphologyEx(logo_edges, cv2.MORPH_CLOSE, kernel)
        logo_edges = cv2.morphologyEx(logo_edges, cv2.MORPH_DILATE, np.ones((2, 2), np.uint8))
        
        logo_contours, _ = cv2.findContours(logo_edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        if logo_contours:
            # 면적이 큰 컨투어들을 로고로 간주 (여러 개 가능)
            for contour in logo_contours:
                area = cv2.contourArea(contour)
                if area > 200:  # 최소 크기 조건 완화
                    x, y, w, h = cv2.boundingRect(contour)
                    # 너무 작거나 선형인 것 제외 (조건 완화)
                    if w > 15 and h > 15 and min(w, h) / max(w, h) > 0.15:
                        # 여백 추가로 더 안전하게 보호
                        margin = 5
                        regions.append({
                            'left': max(0, x - margin),
                            'top': max(0, y - margin),
                            'width': min(w + 2*margin, logo_width - x + margin),
                            'height': min(h + 2*margin, logo_height - y + margin),
                            'type': 'logo'
                        })
                        print(f"  - 📋 로고 영역 감지: {w}x{h} at ({x},{y}) (여백 포함)")
        
        # 2. 도장/인감 영역 감지 (하단 우측) - 빨간 도장
        stamp_start_y = int(height * 0.4)  # 하단 60% 영역
        stamp_start_x = int(width * 0.25)  # 우측 75% 영역
        stamp_region = gray[stamp_start_y:height, stamp_start_x:width]
        
        # 도장 영역에서 원형/사각형 블록 찾기 (개선된 감지)
        # OpenCV Gaussian kernel 오류 방지: ksize 조건 검증
        stamp_ksize = (3, 3)  # (2,2)는 짝수이므로 (3,3)으로 변경
        if stamp_ksize[0] > 0 and stamp_ksize[0] % 2 == 1 and stamp_ksize[1] > 0 and stamp_ksize[1] % 2 == 1:
            stamp_blur = cv2.GaussianBlur(stamp_region, stamp_ksize, 0)
        else:
            stamp_blur = stamp_region.copy()  # 블러 없이 원본 사용
        stamp_edges = cv2.Canny(stamp_blur, 25, 100)
        
        # 모폴로지 연산으로 도장 형태 강화
        kernel_stamp = np.ones((3, 3), np.uint8)
        stamp_edges = cv2.morphologyEx(stamp_edges, cv2.MORPH_CLOSE, kernel_stamp)
        stamp_edges = cv2.morphologyEx(stamp_edges, cv2.MORPH_DILATE, np.ones((2, 2), np.uint8))
        
        stamp_contours, _ = cv2.findContours(stamp_edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        if stamp_contours:
            # 원형/사각형에 가까운 컨투어를 도장으로 간주
            for contour in stamp_contours:
                area = cv2.contourArea(contour)
                if area > 100:  # 최소 크기 조건 완화
                    x, y, w, h = cv2.boundingRect(contour)
                    # 도장 형태 확인 (정사각형 또는 원형) - 조건 완화
                    aspect_ratio = float(w) / h if h > 0 else 1
                    if 0.5 <= aspect_ratio <= 2.0 and w > 12 and h > 12:  # 도장 형태
                        # 여백 추가로 더 안전하게 보호
                        margin = 8
                        regions.append({
                            'left': max(0, x + stamp_start_x - margin),
                            'top': max(0, y + stamp_start_y - margin),
                            'width': min(w + 2*margin, width - (x + stamp_start_x) + margin),
                            'height': min(h + 2*margin, height - (y + stamp_start_y) + margin),
                            'type': 'stamp'
                        })
                        print(f"  - 🔴 도장 영역 감지: {w}x{h} at ({x + stamp_start_x},{y + stamp_start_y}) (여백 포함)")
        
        # 3. 벡터 그래픽 요소 감지 (선, 도형, 표)
        # 수평선 감지 (개선)
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 1))
        horizontal_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, horizontal_kernel)
        horizontal_contours, _ = cv2.findContours(horizontal_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        for contour in horizontal_contours:
            area = cv2.contourArea(contour)
            if area > 50:  # 조건 완화
                x, y, w, h = cv2.boundingRect(contour)
                if w > 30 and h < 15:  # 긴 수평선
                    margin = 2
                    regions.append({
                        'left': max(0, x - margin),
                        'top': max(0, y - margin),
                        'width': min(w + 2*margin, width - x + margin),
                        'height': min(h + 2*margin, height - y + margin),
                        'type': 'line'
                    })
                    print(f"  - ➖ 수평선 감지: {w}x{h} at ({x},{y})")
        
        # 수직선 감지 (개선)
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 30))
        vertical_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, vertical_kernel)
        vertical_contours, _ = cv2.findContours(vertical_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        for contour in vertical_contours:
            area = cv2.contourArea(contour)
            if area > 50:  # 조건 완화
                x, y, w, h = cv2.boundingRect(contour)
                if h > 30 and w < 15:  # 긴 수직선
                    margin = 2
                    regions.append({
                        'left': max(0, x - margin),
                        'top': max(0, y - margin),
                        'width': min(w + 2*margin, width - x + margin),
                        'height': min(h + 2*margin, height - y + margin),
                        'type': 'line'
                    })
                    print(f"  - ⬇️ 수직선 감지: {w}x{h} at ({x},{y})")
        
        # 4. 기타 벡터 요소 감지 (도형, 아이콘 등)
        # 엣지 기반 도형 감지
        edges = cv2.Canny(gray, 20, 100)
        kernel_shape = np.ones((3, 3), np.uint8)
        edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel_shape)
        
        shape_contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        for contour in shape_contours:
            area = cv2.contourArea(contour)
            if 80 <= area <= 5000:  # 중간 크기 도형
                x, y, w, h = cv2.boundingRect(contour)
                aspect_ratio = float(w) / h if h > 0 else 1
                
                # 정사각형, 원형, 직사각형 등의 벡터 요소
                if (0.3 <= aspect_ratio <= 3.0 and w > 8 and h > 8 and 
                    w < width * 0.8 and h < height * 0.8):  # 너무 큰 것 제외
                    
                    # 텍스트 영역이 아닌지 확인 (밀도 체크)
                    roi = gray[y:y+h, x:x+w]
                    if roi.size > 0:
                        # 흰색 픽셀 비율로 텍스트 여부 판단
                        white_ratio = np.sum(roi > 200) / roi.size
                        if white_ratio < 0.7:  # 텍스트가 아닌 그래픽 요소
                            margin = 3
                            regions.append({
                                'left': max(0, x - margin),
                                'top': max(0, y - margin),
                                'width': min(w + 2*margin, width - x + margin),
                                'height': min(h + 2*margin, height - y + margin),
                                'type': 'vector'
                            })
                            print(f"  - 🔷 벡터 요소 감지: {w}x{h} at ({x},{y})")
        
        # 5. 전체 레이아웃을 배경 이미지로 보존 (가장 중요!)
        regions.append({
            'left': 0,
            'top': 0,
            'width': width,
            'height': height,
            'type': 'background'
        })
        print(f"  - 🖼️ 배경 레이아웃 보존: {width}x{height}")
        
        print(f"  - ✅ 총 {len(regions)}개 이미지/벡터 영역 감지됨 (로고, 도장, 선, 벡터, 배경)")
        return regions
        
    except Exception as e:
        print(f"  - ❌ 이미지 영역 감지 오류: {e}")
        # 오류 시에도 전체 레이아웃은 보존
        return [{
            'left': 0,
            'top': 0,
            'width': image.size[0],
            'height': image.size[1],
            'type': 'background'
        }]

def detect_document_type(image, text_blocks):
    """문서 타입 감지 (공문서 특화 - 항상 하이브리드 모드)"""
    try:
        img_width, img_height = image.size
        total_area = img_width * img_height
        
        # 텍스트 영역 비율 계산
        text_area = 0
        for block in text_blocks:
            text_area += block['width'] * block['height']
        
        text_ratio = text_area / total_area if total_area > 0 else 0
        
        print(f"  - 텍스트 영역 비율: {text_ratio:.2%}")
        print(f"  - 텍스트 블록 수: {len(text_blocks)}개")
        
        # 공문서는 항상 하이브리드 모드로 처리 (최고 품질)
        doc_type = "hybrid"
        print("  - 📄 공문서 특화 모드: 하이브리드 변환 (원본 레이아웃 완벽 보존 + 완전 편집 가능)")
                
        return doc_type
        
    except Exception as e:
        print(f"문서 타입 감지 오류: {e}")
        return "hybrid"  # 기본값

def add_text_only_conversion(doc, image, section, text_blocks):
    """공문서 텍스트 전용 변환: 완벽한 서식 보존 및 편집 기능"""
    try:
        print("  - 📄 공문서 텍스트 전용 변환 모드 시작 (완벽한 서식 보존)")
        
        if not text_blocks:
            print("  - ⚠️ 텍스트 블록이 없어 빈 문서 생성")
            para = doc.add_paragraph("텍스트를 추출할 수 없었습니다.")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        
        # 텍스트 블록을 Y 좌표 기준으로 정렬
        sorted_blocks = sorted(text_blocks, key=lambda x: (x['top'], x['left']))
        
        # 이미지 영역 감지 (텍스트와 충돌 방지용)
        image_regions = detect_image_regions(image)
        
        # 겹침 방지 및 이미지 영역 충돌 회피 적용
        adjusted_blocks = _prevent_text_overlap(sorted_blocks, image_regions)
        
        print(f"  - 📄 {len(adjusted_blocks)}개 텍스트 블록 처리 (공문서 특화 서식 적용)")
        
        # 공문서 특화 텍스트 배치
        for i, block in enumerate(adjusted_blocks):
            para = doc.add_paragraph()
            run = para.add_run(block['text'])
            
            # 공문서 표준 폰트 설정
            run.font.name = "맑은 고딕"
            run.font.size = Pt(11)  # 공문서 표준 크기
            
            # 신뢰도에 따른 색상 (편집 시 참고용)
            if block.get('confidence', 85) < 70:
                run.font.color.rgb = RGBColor(64, 64, 64)  # 회색 (낮은 신뢰도)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)  # 검정 (높은 신뢰도)
            
            # 원본 위치를 정확히 반영한 들여쓰기 계산
            img_width = image.size[0]
            page_width_pt = section.page_width.pt - section.left_margin.pt - section.right_margin.pt
            
            # X 좌표 비율을 페이지 너비에 적용
            left_ratio = block['left'] / img_width if img_width > 0 else 0
            left_indent = Pt(left_ratio * page_width_pt * 0.9)  # 90% 비율로 정확한 위치
            
            # 공문서 특화 문단 서식 설정
            para.paragraph_format.left_indent = left_indent
            para.paragraph_format.line_spacing = 1.2  # 공문서 표준 줄간격
            para.paragraph_format.space_after = Pt(2)  # 문단 간격
            
            # Y 좌표를 고려한 상단 여백 설정
            if i > 0:
                prev_block = adjusted_blocks[i-1]
                y_distance = block['top'] - (prev_block['top'] + prev_block['height'])
                if y_distance > 15:  # 충분한 간격이 있는 경우
                    para.paragraph_format.space_before = Pt(min(y_distance * 0.08, 8))
            
            # 텍스트 정렬 (공문서 특화)
            text_content = block['text'].strip()
            
            # 제목이나 중요 텍스트 감지 (중앙 정렬)
            if (left_ratio > 0.3 and left_ratio < 0.7) or any(keyword in text_content for keyword in ['공문', '통지', '안내', '요청', '회신']):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.bold = True  # 제목은 굵게
            elif left_ratio > 0.7:  # 우측에 위치 (날짜, 서명 등)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:  # 좌측에 위치 (본문)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 공문서 특수 서식 감지 및 적용
            if '기관명' in text_content or '부서명' in text_content:
                run.font.size = Pt(12)  # 기관명은 조금 크게
                run.font.bold = True
            elif any(keyword in text_content for keyword in ['수신', '발신', '제목', '내용']):
                run.font.bold = True  # 항목명은 굵게
            
            # 편집 가능성을 위한 스타일 설정
            para.style = doc.styles['Normal']  # 기본 스타일 사용으로 완전 편집 가능
            
            print(f"    텍스트 {i+1}: '{block['text'][:25]}...' (위치: {left_indent.pt:.1f}pt, 신뢰도: {block.get('confidence', 85)}%)")
        
        print(f"  - ✅ 공문서 텍스트 전용 변환 완료: {len(adjusted_blocks)}개 블록 (완벽한 서식 보존 + 완전 편집 가능)")
        return True
        
    except Exception as e:
        print(f"  - ❌ 텍스트 전용 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

def add_hybrid_conversion(doc, image, section, text_blocks):
    """공문서 하이브리드 변환: 원본 레이아웃 완벽 보존 + 완전 편집 가능"""
    try:
        print("  - 📄 공문서 하이브리드 변환 모드 시작 (원본 레이아웃 완벽 보존 + 완전 편집 가능)")
        
        # 본문 영역 크기 계산
        max_w_in = float(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
        max_h_in = float(section.page_height.inches - (section.top_margin.inches + section.bottom_margin.inches))
        
        dpi = 200
        img_w_in = image.size[0] / dpi
        img_h_in = image.size[1] / dpi
        fit_w, fit_h = _fit_dimensions_within(max_w_in, max_h_in, img_w_in, img_h_in)
        
        # 1. 배경 이미지 추가 (원본 레이아웃 보존)
        print("  - 🖼️ 배경 이미지 추가 (원본 레이아웃 보존)")
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                temp_path = tmp.name
                image.save(temp_path, 'JPEG', quality=95, optimize=True)
            
            # 배경 이미지를 문서에 추가
            doc.add_picture(temp_path, width=fit_w, height=fit_h)
            
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
        
        # 2. 편집 가능한 텍스트 블록 추가
        if text_blocks:
            print(f"  - ✏️ 편집 가능한 텍스트 블록 {len(text_blocks)}개 추가")
            
            # 텍스트 블록을 Y 좌표 기준으로 정렬
            sorted_blocks = sorted(text_blocks, key=lambda x: (x['top'], x['left']))
            
            # 이미지 영역 감지 (텍스트와 충돌 방지용)
            image_regions = detect_image_regions(image)
            
            # 겹침 방지 및 이미지 영역 충돌 회피 적용
            adjusted_blocks = _prevent_text_overlap(sorted_blocks, image_regions)
            
            # 각 텍스트 블록을 편집 가능한 문단으로 추가
            for i, block in enumerate(adjusted_blocks):
                para = doc.add_paragraph()
                run = para.add_run(block['text'])
                
                # 한글 폰트 설정
                run.font.name = "맑은 고딕"
                run.font.size = Pt(11)  # 공문서 표준 크기
                
                # 신뢰도에 따른 색상 (편집 시 참고용)
                if block.get('confidence', 85) < 70:
                    run.font.color.rgb = RGBColor(64, 64, 64)  # 회색 (낮은 신뢰도)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 검정 (높은 신뢰도)
                
                # 정확한 위치 매핑을 위한 고정밀 스케일링 계산
                img_width, img_height = image.size
                page_width_pt = section.page_width.pt - section.left_margin.pt - section.right_margin.pt
                page_height_pt = section.page_height.pt - section.top_margin.pt - section.bottom_margin.pt
                
                # 이미지가 페이지에 맞춰진 실제 크기 계산 (고정밀)
                dpi = 200
                img_w_in = img_width / dpi
                img_h_in = img_height / dpi
                max_w_in = page_width_pt / 72.0  # pt to inch (정확한 변환)
                max_h_in = page_height_pt / 72.0  # pt to inch (정확한 변환)
                
                # 비율 유지하며 맞춤 (고정밀 계산)
                scale_x = max_w_in / img_w_in
                scale_y = max_h_in / img_h_in
                scale = min(scale_x, scale_y)  # 비율 유지를 위한 최소값 사용
                
                actual_img_width_pt = img_w_in * scale * 72.0  # inch to pt (정확한 변환)
                actual_img_height_pt = img_h_in * scale * 72.0  # inch to pt (정확한 변환)
                
                # 정확한 X, Y 좌표 매핑 (픽셀 단위 정밀도)
                x_scale = actual_img_width_pt / img_width
                y_scale = actual_img_height_pt / img_height
                
                # 텍스트 블록의 실제 위치 계산 (픽셀 → pt 고정밀 변환)
                actual_left_pt = block['left'] * x_scale
                actual_top_pt = block['top'] * y_scale
                
                # 위치 정확도 보정 (공문서 특화)
                # 좌측 여백 보정 (공문서 표준 여백 고려)
                left_margin_correction = 5.0  # pt 단위 미세 조정
                actual_left_pt = max(0, actual_left_pt - left_margin_correction)
                
                # 들여쓰기 설정 (페이지 여백 고려)
                left_indent = Pt(min(actual_left_pt, page_width_pt * 0.95))  # 최대 95%까지만
                
                # 문단 서식 설정 (공문서 특화)
                para.paragraph_format.left_indent = left_indent
                para.paragraph_format.line_spacing = 1.2  # 공문서 표준 줄간격
                para.paragraph_format.space_after = Pt(2)  # 문단 간격
                
                # Y 좌표를 고려한 정확한 상단 여백 설정
                if i > 0:
                    prev_block = adjusted_blocks[i-1]
                    prev_actual_top = prev_block['top'] * y_scale
                    prev_actual_height = prev_block['height'] * y_scale
                    current_actual_top = actual_top_pt
                    
                    # 실제 픽셀 거리를 pt로 변환
                    y_distance_pt = current_actual_top - (prev_actual_top + prev_actual_height)
                    if y_distance_pt > 10:  # 10pt 이상 간격이 있는 경우
                        para.paragraph_format.space_before = Pt(min(y_distance_pt * 0.6, 15))  # 최대 15pt
                
                # 정확한 위치 기반 텍스트 정렬 결정
                left_ratio = actual_left_pt / page_width_pt if page_width_pt > 0 else 0
                text_width_pt = len(block['text']) * 6  # 대략적인 텍스트 폭 추정
                text_center_ratio = (actual_left_pt + text_width_pt/2) / page_width_pt
                
                # 더 정확한 정렬 판단
                if 0.4 <= text_center_ratio <= 0.6:  # 텍스트 중심이 페이지 중앙 부근
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif left_ratio > 0.75:  # 우측에 위치 (날짜, 서명 등)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:  # 좌측에 위치 (본문)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 공문서 특수 텍스트 감지 및 서식 적용
                text_content = block['text'].strip()
                if any(keyword in text_content for keyword in ['공문', '통지', '안내', '요청', '제목']):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.bold = True
                    run.font.size = Pt(12)
                elif any(keyword in text_content for keyword in ['수신', '발신', '내용']):
                    run.font.bold = True
                
                print(f"    텍스트 {i+1}: '{block['text'][:25]}...' (위치: {left_indent.pt:.1f}pt, Y: {actual_top_pt:.1f}pt, 신뢰도: {block.get('confidence', 85)}%)")
        
        print(f"  - ✅ 공문서 하이브리드 변환 완료: 원본 레이아웃 완벽 보존 + {len(text_blocks)}개 텍스트 블록 편집 가능")
        return True
        
    except Exception as e:
        print(f"  - ❌ 하이브리드 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

def add_image_only_conversion(doc, image, section):
    """4번 이미지: 이미지 블럭 그대로 유지 모드"""
    try:
        print("  - 🖼️ 이미지 전용 변환 모드 시작")
        
        # 본문 영역 크기 계산
        max_w_in = float(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
        max_h_in = float(section.page_height.inches - (section.top_margin.inches + section.bottom_margin.inches))
        
        dpi = 200
        img_w_in = image.size[0] / dpi
        img_h_in = image.size[1] / dpi
        fit_w, fit_h = _fit_dimensions_within(max_w_in, max_h_in, img_w_in, img_h_in)
        
        # 원본 이미지만 추가 (편집 불가)
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                temp_path = tmp.name
                image.save(temp_path, 'JPEG', quality=90, optimize=True)
            
            # 이미지를 직접 문서에 추가
            doc.add_picture(temp_path, width=fit_w, height=fit_h)
            print("  - ✅ 이미지 전용 변환 완료: 원본 이미지만 유지")
            
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
                    
        return True
        
    except Exception as e:
        print(f"  - ❌ 이미지 전용 변환 오류: {e}")
        return False

def _add_textbox_paragraph(doc, left_pt: float, top_pt: float, width_pt: float, height_pt: float, text: str):
    """완전 편집 가능한 텍스트 문단 추가 (공문서 특화)"""
    try:
        # 텍스트 정리
        clean_text = (text or "").strip()
        if not clean_text:
            return True
        
        # 문단 생성
        para = doc.add_paragraph()
        run = para.add_run(clean_text)
        
        # 공문서 표준 폰트 설정
        run.font.name = "맑은 고딕"
        run.font.size = Pt(11)  # 공문서 표준 크기
        run.font.color.rgb = RGBColor(0, 0, 0)  # 검정색
        
        # 위치 기반 들여쓰기 계산 (left_pt를 기준으로)
        # A4 페이지 기준 (595pt 너비)
        page_width_pt = 595
        left_ratio = left_pt / page_width_pt if page_width_pt > 0 else 0
        left_indent = Pt(left_ratio * page_width_pt * 0.85)  # 85% 비율로 정확한 위치
        
        # 문단 서식 설정 (공문서 특화)
        para.paragraph_format.left_indent = left_indent
        para.paragraph_format.line_spacing = 1.2  # 공문서 표준 줄간격
        para.paragraph_format.space_after = Pt(2)  # 문단 간격
        
        # 텍스트 정렬 (위치에 따른 자동 정렬)
        if left_ratio > 0.3 and left_ratio < 0.7:  # 중앙 부근
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif left_ratio > 0.7:  # 우측
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:  # 좌측
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Y 위치에 따른 상단 여백 설정
        if top_pt > 50:  # 페이지 상단이 아닌 경우
            para.paragraph_format.space_before = Pt(min(top_pt * 0.02, 8))
        
        # 편집 가능성을 위한 스타일 설정
        para.style = doc.styles['Normal']  # 기본 스타일 사용으로 완전 편집 가능
        
        return True
        
    except Exception as e:
        print(f"텍스트 문단 추가 오류: {e}")
        return False

def docx_to_pdf(docx_path, output_path):
    """DOCX를 PDF로 변환 (한글 폰트 지원)"""
    try:
        print(f"DOCX → PDF 변환 시작: {docx_path}")
        
        # 한글 폰트 설정
        font_setup = setup_korean_fonts()
        
        # DOCX 문서 읽기
        doc = Document(docx_path)
        
        # PDF 생성
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        y_position = height - 50
        
        # 한글 폰트 사용
        if font_setup:
            c.setFont("Korean", 12)
        else:
            c.setFont("Helvetica", 12)
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 텍스트 처리 (한글 지원)
                text = paragraph.text.strip()
                
                # 긴 텍스트를 여러 줄로 나누기
                max_chars_per_line = 80
                lines = []
                while len(text) > max_chars_per_line:
                    lines.append(text[:max_chars_per_line])
                    text = text[max_chars_per_line:]
                if text:
                    lines.append(text)
                
                # 각 줄을 PDF에 추가
                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                        if font_setup:
                            c.setFont("Korean", 12)
                        else:
                            c.setFont("Helvetica", 12)
                    
                    c.drawString(50, y_position, line)
                    y_position -= 20
        
        c.save()
        print(f"✅ DOCX → PDF 변환 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ DOCX → PDF 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

@app.route('/')
def index():
    return render_template('index.html')

# 정적 파일 서빙
@app.route('/static/<path:filename>')
def static_files(filename):
    return send_file(os.path.join('static', filename))

@app.route('/convert', methods=['POST'])
def convert_file():
    try:
        print("=== 변환 요청 시작 ===")
        
        # 1. 파일 확인
        if 'file' not in request.files:
            print("오류: 파일이 없음")
            return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
        
        file = request.files['file']
        if file.filename == '':
            print("오류: 파일명이 없음")
            return jsonify({'success': False, 'error': '파일명이 비어있습니다.'}), 400
        
        # 2. 파일 형식 확인
        if not allowed_file(file.filename, file.content_type):
            print(f"오류: 지원하지 않는 파일 형식 - {file.filename} (MIME: {file.content_type})")
            return jsonify({'success': False, 'error': 'PDF 또는 DOCX 파일만 업로드 가능합니다.'}), 400
        
        # 3. 파일명 처리 및 저장
        original_filename = file.filename
        print(f"원본 파일명: {original_filename}")
        
        # 파일명 정리 및 처리
        cleaned_filename = original_filename.strip()
        if not cleaned_filename:
            cleaned_filename = "uploaded_file"
        
        # 파일명과 확장자를 분리하여 안전하게 처리
        if '.' in cleaned_filename:
            # 확장자가 있는 경우
            name_part, ext_part = cleaned_filename.rsplit('.', 1)
            safe_name = secure_filename(name_part) or "file"
            safe_ext = ext_part.lower().strip()
            
            # 확장자가 비어있으면 기본값 설정
            if not safe_ext:
                safe_ext = "pdf"
            
            filename = f"{safe_name}.{safe_ext}"
            print(f"확장자 분리 처리: {cleaned_filename} → {filename}")
        else:
            # 확장자가 없는 경우 MIME 타입으로 추정
            content_type = file.content_type
            print(f"파일 MIME 타입: {content_type}")
            
            safe_name = secure_filename(cleaned_filename) or "file"
            
            if 'pdf' in content_type:
                filename = f"{safe_name}.pdf"
                print(f"PDF 파일로 추정하여 .pdf 확장자 추가")
            elif 'document' in content_type or 'word' in content_type:
                filename = f"{safe_name}.docx"
                print(f"DOCX 파일로 추정하여 .docx 확장자 추가")
            else:
                # MIME 타입도 없는 경우 기본값으로 처리
                print(f"경고: 파일에 확장자와 MIME 타입이 없음 - {cleaned_filename}")
                filename = f"{safe_name}.pdf"
                print(f"기본값으로 .pdf 확장자 추가")
        
        # 최종 파일명 검증 및 보정
        if not filename or filename == '.' or filename == '..' or '.' not in filename:
            filename = "uploaded_file.pdf"
            print(f"안전하지 않은 파일명으로 인해 기본 파일명 사용: {filename}")
        
        # 확장자 최종 검증
        if not filename.endswith(('.pdf', '.docx')):
            if filename.endswith('.pdf') or 'pdf' in file.content_type:
                filename = filename.rsplit('.', 1)[0] + '.pdf'
            else:
                filename = filename.rsplit('.', 1)[0] + '.docx'
            print(f"확장자 보정: {filename}")
        
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        print(f"파일 저장 완료: {file_path}")
        
        # 4. 파일 확장자 확인 (간단하고 안전한 처리)
        print(f"최종 파일명: {filename}")
        
        # 확장자 추출 (이미 위에서 안전하게 처리되었으므로 간단하게)
        if '.' not in filename:
            print(f"오류: 최종 파일명에 확장자가 없음 - {filename}")
            return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
        
        try:
            file_ext = filename.split('.')[-1].lower().strip()
            if not file_ext:
                print(f"오류: 확장자가 비어있음 - {filename}")
                return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
            
            print(f"파일 확장자: {file_ext}")
        except Exception as e:
            print(f"오류: 파일 확장자 추출 중 예외 발생 - {e}")
            return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
        
        # 5. 변환 처리
        try:
            # 출력 파일명 생성 (안전하게)
            base_name = filename.rsplit('.', 1)[0].strip()
            if not base_name:
                base_name = "converted_file"  # 기본 파일명
            
            if file_ext == 'pdf':
                # PDF → DOCX
                output_filename = base_name + '.docx'
                output_path = os.path.join(OUTPUT_FOLDER, output_filename)
                print(f"PDF → DOCX 변환: {file_path} → {output_path}")
                success = pdf_to_docx(file_path, output_path)
                
            elif file_ext == 'docx':
                # DOCX → PDF
                output_filename = base_name + '.pdf'
                output_path = os.path.join(OUTPUT_FOLDER, output_filename)
                print(f"DOCX → PDF 변환: {file_path} → {output_path}")
                success = docx_to_pdf(file_path, output_path)
            else:
                print(f"오류: 지원하지 않는 파일 형식 - {file_ext}")
                return jsonify({'success': False, 'error': f'지원하지 않는 파일 형식입니다: {file_ext}'}), 400
                
        except Exception as e:
            print(f"오류: 변환 처리 중 예외 발생 - {e}")
            return jsonify({'success': False, 'error': '파일 변환 처리 중 오류가 발생했습니다.'}), 500
        
        # 6. 임시 파일 정리 (안전한 삭제)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print("✅ 임시 파일 삭제 완료")
        except Exception as cleanup_error:
            print(f"⚠️ 임시 파일 삭제 실패: {cleanup_error}")
        
        # 7. 결과 처리 및 검증
        if success and os.path.exists(output_path):
            # 출력 파일 크기 확인
            try:
                output_size = os.path.getsize(output_path) / 1024  # KB
                print(f"📄 변환된 파일 크기: {output_size:.1f}KB")
                
                if output_size < 1:  # 1KB 미만
                    print("⚠️ 변환된 파일이 너무 작습니다 - 변환 품질 확인 필요")
                
                print(f"✅ 변환 성공! 다운로드: {output_filename}")
                return send_file(output_path, as_attachment=True, download_name=output_filename)
                
            except Exception as size_check_error:
                print(f"⚠️ 파일 크기 확인 실패: {size_check_error}")
                print(f"✅ 변환 성공! 다운로드: {output_filename}")
                return send_file(output_path, as_attachment=True, download_name=output_filename)
        else:
            print("❌ 변환 실패")
            error_msg = "파일 변환에 실패했습니다."
            if not success:
                error_msg += " (변환 프로세스 오류)"
            if not os.path.exists(output_path):
                error_msg += " (출력 파일 생성 실패)"
            return jsonify({'success': False, 'error': error_msg}), 500
            
    except Exception as e:
        print(f"❌ 서버 오류: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'서버 오류: {str(e)}'}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({'success': False, 'error': '파일 크기가 100MB를 초과합니다.'}), 413

if __name__ == '__main__':
    print("🚀 PDF ↔ DOCX 변환기 시작")
    print("📍 서버 주소: http://127.0.0.1:5000")
    print("📍 네트워크: http://0.0.0.0:5000")
    
    # 한글 폰트 설정
    print("🔤 한글 폰트 설정 중...")
    font_setup = setup_korean_fonts()
    if font_setup:
        print("✅ 한글 폰트 설정 완료")
    else:
        print("⚠️ 한글 폰트 설정 실패 - 기본 폰트 사용")
    
    app.run(debug=True, host='0.0.0.0', port=5000)