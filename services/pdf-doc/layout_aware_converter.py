import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import io
import os
import logging
import math

# 로깅 설정
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class LayoutAwareConverter: 
     """ 
     PDF 요소의 위치, 크기, 정렬을 분석하여 원본 레이아웃을 최대한 보존하는 
     지능형 변환기입니다. 
     """
     
     def __init__(self):
         self.page_width = 0
         self.page_height = 0
         self.scale_factor = 1.0 
 
     def _get_alignment(self, block_bbox, page_rect): 
         """텍스트 블록의 위치를 기반으로 정렬을 결정합니다.""" 
         x0, _, x1, _ = block_bbox 
         page_width = page_rect.width 
         center_x = page_width / 2 
         block_center_x = (x0 + x1) / 2 
         block_width = x1 - x0
         
         # 더 정확한 정렬 판단을 위한 개선된 로직
         left_margin = x0 / page_width
         right_margin = (page_width - x1) / page_width
         center_offset = abs(block_center_x - center_x) / page_width
         
         # 중앙 정렬: 좌우 여백이 비슷하고 중앙에서 크게 벗어나지 않음
         if center_offset < 0.05 and abs(left_margin - right_margin) < 0.1:
             return WD_ALIGN_PARAGRAPH.CENTER
         # 우측 정렬: 우측 여백이 작고 좌측 여백이 큼
         elif right_margin < 0.15 and left_margin > 0.3:
             return WD_ALIGN_PARAGRAPH.RIGHT
         # 기본값은 좌측 정렬
         else:
             return WD_ALIGN_PARAGRAPH.LEFT
             
     def _calculate_spacing(self, current_bbox, previous_bbox, page_rect):
         """두 요소 간의 간격을 계산하여 적절한 여백을 반환합니다."""
         if previous_bbox is None:
             return 0
             
         # 수직 간격 계산 (포인트 단위를 cm로 변환)
         vertical_gap = current_bbox[1] - previous_bbox[3]
         vertical_gap_cm = max(0, vertical_gap * 0.0352778)  # pt to cm
         
         # 간격이 너무 작으면 최소값 적용, 너무 크면 최대값 적용
         return min(max(vertical_gap_cm, 0.1), 2.0)
         
     def _add_spacing_before_paragraph(self, paragraph, spacing_cm):
         """단락 앞에 정확한 간격을 추가합니다."""
         if spacing_cm > 0.2:  # 0.2cm 이상일 때만 간격 추가
             paragraph.paragraph_format.space_before = Cm(spacing_cm) 
 
     def _process_image(self, img_data): 
         """이미지를 검증하고 DOCX 삽입에 안전한 형식으로 처리합니다.""" 
         try: 
             img = Image.open(io.BytesIO(img_data)) 
             if img.mode in ('RGBA', 'LA', 'P'): 
                 background = Image.new("RGB", img.size, (255, 255, 255)) 
                 img_rgba = img.convert("RGBA") 
                 background.paste(img_rgba, mask=img_rgba.split()[-1]) 
                 img = background 
             elif img.mode == 'CMYK': 
                 img = img.convert("RGB") 
             
             output_stream = io.BytesIO() 
             img.save(output_stream, format="PNG") 
             return output_stream.getvalue() 
         except Exception as e: 
             logging.error(f"이미지 처리 실패: {e}") 
             return None 
 
     def convert(self, pdf_path, docx_path): 
         """PDF를 레이아웃을 보존하여 DOCX로 변환합니다.""" 
         logging.info(f"'{pdf_path}' 파일 변환 시작...") 
         pdf_doc = fitz.open(pdf_path) 
         docx_doc = Document() 
         
         # A4 용지 크기에 맞춰 여백 설정 (선택 사항) 
         section = docx_doc.sections[0] 
         section.left_margin = Cm(2.5) 
         section.right_margin = Cm(2.5) 
         section.top_margin = Cm(2.5) 
         section.bottom_margin = Cm(2.5) 
 
         for page_num, page in enumerate(pdf_doc): 
             logging.info(f"--- 페이지 {page_num + 1} 처리 중 ---") 
             
             # 페이지 크기 정보 저장
             self.page_width = page.rect.width
             self.page_height = page.rect.height
             
             # 1. 페이지의 모든 요소(텍스트, 이미지)를 위치 정보와 함께 추출 
             elements = []
             previous_bbox = None 
             
             # 텍스트 블록 추출 
             text_blocks = page.get_text("dict")["blocks"] 
             for block in text_blocks: 
                 if block['type'] == 0:  # 0은 텍스트 블록 
                     elements.append({ 
                         "type": "text", 
                         "bbox": block['bbox'], 
                         "content": block 
                     }) 
             
             # 이미지 블록 추출 
             images = page.get_images(full=True) 
             for img_info in images: 
                 xref = img_info[0] 
                 try: 
                     img_bbox = page.get_image_bbox(img_info) 
                     elements.append({ 
                         "type": "image", 
                         "bbox": img_bbox, 
                         "content": xref 
                     }) 
                 except ValueError: 
                     logging.warning(f"이미지 바운딩 박스 추출 실패: xref {xref}") 
 
 
             # 2. 모든 요소를 수직(y) > 수평(x) 순서로 정렬 
             elements.sort(key=lambda el: (el['bbox'][1], el['bbox'][0])) 
 
             # 3. 정렬된 순서대로 DOCX에 요소 추가 
             for elem in elements: 
                 if elem['type'] == 'text': 
                     block = elem['content'] 
                     p = docx_doc.add_paragraph() 
                     
                     # 정렬 설정 
                     p.alignment = self._get_alignment(block['bbox'], page.rect) 
                     
                     # 이전 요소와의 간격 계산 및 적용
                     spacing = self._calculate_spacing(block['bbox'], previous_bbox, page.rect)
                     self._add_spacing_before_paragraph(p, spacing)
                     
                     # 줄 간격 설정 (원본과 유사하게)
                     p.paragraph_format.line_spacing = 1.0
                     p.paragraph_format.space_after = Pt(0)
 
                     # 텍스트 내용 추가
                     text_content = ""
                     for line in block['lines']: 
                         for span in line['spans']: 
                             text_content += span['text']
                     
                     if text_content.strip():  # 빈 텍스트가 아닌 경우만 처리
                         run = p.add_run(text_content)
                         
                         # 첫 번째 span의 속성을 기준으로 폰트 설정
                         if block['lines'] and block['lines'][0]['spans']:
                             first_span = block['lines'][0]['spans'][0]
                             run.font.size = Pt(max(8, int(first_span['size'])))
                             
                             # 굵은 글씨, 이탤릭체 감지 
                             if "bold" in first_span['font'].lower(): 
                                 run.bold = True 
                             if "italic" in first_span['font'].lower(): 
                                 run.italic = True
                     
                     previous_bbox = block['bbox'] 
                 
                 elif elem['type'] == 'image': 
                     xref = elem['content'] 
                     try: 
                         pix = fitz.Pixmap(pdf_doc, xref) 
                         img_data = pix.tobytes("png") 
                         
                         processed_data = self._process_image(img_data) 
                         if not processed_data: 
                             continue 
 
                         # 이미지 크기 계산 (PDF에서의 크기를 DOCX 단위로 변환) 
                         img_width_pt = elem['bbox'][2] - elem['bbox'][0]
                         img_height_pt = elem['bbox'][3] - elem['bbox'][1]

                         # 너비와 높이가 음수가 되지 않도록 보장
                         img_width_cm = max(1.0, img_width_pt * 0.0352778)  # 최소 1cm
                         img_height_cm = max(1.0, img_height_pt * 0.0352778)  # 최소 1cm
                         
                         # 페이지 너비를 초과하지 않도록 제한 (A4 기준 약 16cm)
                         max_width_cm = 16.0
                         if img_width_cm > max_width_cm:
                             ratio = max_width_cm / img_width_cm
                             img_width_cm = max_width_cm
                             img_height_cm *= ratio
                         
                         p = docx_doc.add_paragraph() 
                         
                         # 이미지 위치에 따른 정렬 결정
                         img_alignment = self._get_alignment(elem['bbox'], page.rect)
                         p.alignment = img_alignment
                         
                         # 이전 요소와의 간격 계산 및 적용
                         spacing = self._calculate_spacing(elem['bbox'], previous_bbox, page.rect)
                         self._add_spacing_before_paragraph(p, spacing)
                         
                         run = p.add_run() 
                         run.add_picture(io.BytesIO(processed_data), width=Cm(img_width_cm), height=Cm(img_height_cm)) 
                         logging.info(f"이미지 추가 완료 (폭: {img_width_cm:.2f}cm, 높이: {img_height_cm:.2f}cm)") 
                         
                         previous_bbox = elem['bbox']
 
                     except Exception as e:
                         logging.error(f"이미지 변환 중 오류 발생: {e}")