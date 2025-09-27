from flask import Flask, render_template, request, jsonify, send_file
import os
import time
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape, portrait
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.utils import ImageReader
import urllib.request
import PyPDF2
import unicodedata
import sys
from PIL import Image as PILImage
import io

# OCR 기능 확인 및 설정
try:
    import pytesseract
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
    print("✅ OCR 모듈 로드 성공")
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR 모듈 없음")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

# 폴더 생성
os.makedirs('uploads', exist_ok=True)
os.makedirs('outputs', exist_ok=True)
os.makedirs('fonts', exist_ok=True)

# 한글 폰트 설정
KOREAN_FONT = 'Helvetica'
KOREAN_FONT_AVAILABLE = False
AVAILABLE_FONTS = {}  # 추가된 변수 정의

def setup_korean_font_advanced():
    """고급 한글 폰트 설정"""
    global KOREAN_FONT, KOREAN_FONT_AVAILABLE, AVAILABLE_FONTS
    
    # 1. 나눔고딕 TTF 시도
    try:
        font_path = os.path.join('fonts', 'NanumGothic.ttf')
        
        if not os.path.exists(font_path):
            print("📥 나눔고딕 폰트 다운로드 중...")
            font_url = "https://github.com/naver/nanumfont/raw/master/TTF/NanumGothic.ttf"
            
            req = urllib.request.Request(font_url, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            
            with urllib.request.urlopen(req, timeout=30) as response:
                with open(font_path, 'wb') as f:
                    f.write(response.read())
        
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
            KOREAN_FONT = 'NanumGothic'
            KOREAN_FONT_AVAILABLE = True
            AVAILABLE_FONTS['NanumGothic'] = {
                'path': font_path,
                'display_name': '나눔고딕'
            }
            print("✅ 나눔고딕 TTF 폰트 등록 완료")
            return True
            
    except Exception as e:
        print(f"나눔고딕 TTF 등록 실패: {e}")
    
    # 2. 시스템 한글 폰트 시도
    system_fonts = [
        (r'C:\Windows\Fonts\malgun.ttf', 'Malgun', '맑은 고딕'),
        (r'C:\Windows\Fonts\gulim.ttc', 'Gulim', '굴림'),
        (r'C:\Windows\Fonts\batang.ttc', 'Batang', '바탕'),
    ]
    
    for font_path, font_name, display_name in system_fonts:
        try:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                KOREAN_FONT = font_name
                KOREAN_FONT_AVAILABLE = True
                AVAILABLE_FONTS[font_name] = {
                    'path': font_path,
                    'display_name': display_name
                }
                print(f"✅ 시스템 한글 폰트 등록: {display_name}")
                return True
        except Exception as e:
            continue
    
    print("⚠️ 한글 폰트 등록 실패, 기본 폰트 사용")
    return False

# 앱 시작 시 한글 폰트 설정
# 기존 setup_korean_font_advanced() 함수를 setup_korean_font_simple()로 교체
# 온라인 폰트 다운로드 코드 제거
# 시스템 폰트만 사용
setup_korean_font_advanced()

def safe_korean_text(text):
    """한글 텍스트 안전 처리"""
    if not text:
        return ""
    
    try:
        normalized = unicodedata.normalize('NFC', str(text))
        cleaned = ''.join(char for char in normalized 
                         if unicodedata.category(char) not in ['Cc', 'Cf'])
        return cleaned if cleaned.strip() else ""
    except Exception as e:
        return str(text) if text else ""

def draw_korean_text(canvas_obj, x, y, text, font_size=11):
    """한글 텍스트 그리기"""
    if not text or not text.strip():
        return
    
    try:
        safe_text = safe_korean_text(text)
        if not safe_text:
            return
        
        if KOREAN_FONT_AVAILABLE:
            canvas_obj.setFont(KOREAN_FONT, font_size)
            canvas_obj.drawString(x, y, safe_text)
        else:
            canvas_obj.setFont('Helvetica', font_size)
            # 한글이 있으면 대체 문자 사용
            has_korean = any('\uac00' <= char <= '\ud7af' for char in safe_text)
            if has_korean:
                display_text = ''.join('한' if '\uac00' <= char <= '\ud7af' else char for char in safe_text)
            else:
                display_text = safe_text
            canvas_obj.drawString(x, y, display_text)
            
    except Exception as e:
        try:
            canvas_obj.setFont('Helvetica', 8)
            canvas_obj.drawString(x, y, "[Error]")
        except:
            pass

def extract_images_from_docx(docx_path, temp_files):
    """DOCX에서 모든 이미지 추출 (강화된 버전)"""
    images = []
    
    try:
        print("🖼️ DOCX에서 이미지 추출 시작...")
        
        # DOCX 파일을 ZIP으로 열어서 이미지 직접 추출
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # media 폴더의 모든 이미지 파일 찾기
            media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
            
            for i, media_file in enumerate(media_files):
                try:
                    # 이미지 파일 확장자 확인
                    if any(media_file.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                        # 이미지 데이터 추출
                        image_data = docx_zip.read(media_file)
                        
                        # PIL로 이미지 정보 확인
                        pil_image = PILImage.open(io.BytesIO(image_data))
                        width, height = pil_image.size
                        
                        # 임시 파일로 저장
                        timestamp = str(int(time.time() * 1000))
                        temp_img_path = os.path.join('uploads', f'extracted_img_{timestamp}_{i}.jpg')
                        
                        # JPEG로 변환하여 저장
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            # 투명도가 있는 이미지는 흰 배경으로 변환
                            background = PILImage.new('RGB', pil_image.size, (255, 255, 255))
                            if pil_image.mode == 'P':
                                pil_image = pil_image.convert('RGBA')
                            background.paste(pil_image, mask=pil_image.split()[-1] if pil_image.mode == 'RGBA' else None)
                            pil_image = background
                        
                        pil_image.save(temp_img_path, 'JPEG', quality=90)
                        temp_files.append(temp_img_path)
                        
                        images.append({
                            'path': temp_img_path,
                            'width': width,
                            'height': height,
                            'original_name': media_file
                        })
                        
                        print(f"✅ 이미지 추출: {media_file} ({width}x{height})")
                        
                except Exception as e:
                    print(f"이미지 {media_file} 추출 오류: {e}")
                    continue
        
        print(f"✅ 총 {len(images)}개 이미지 추출 완료")
        return images
        
    except Exception as e:
        print(f"❌ 이미지 추출 실패: {e}")
        return []

import zipfile

def extract_docx_with_complete_formatting(docx_path, temp_files):
    """DOCX에서 완전한 서식 정보와 함께 내용 추출"""
    try:
        doc = Document(docx_path)
        all_content = []
        
        print("📝 완전한 서식 정보와 함께 내용 추출 시작...")
        
        # 1. 이미지 먼저 추출
        extracted_images = extract_images_from_docx(docx_path, temp_files)
        image_index = 0
        
        # 2. 문단별 서식 정보 추출 (이미지 포함)
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                # 문단에 이미지가 있는지 확인
                has_image = False
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        # drawing 요소 확인 (이미지)
                        drawings = run._element.xpath('.//a:blip')
                        if drawings and image_index < len(extracted_images):
                            # 이미지 추가
                            img_info = extracted_images[image_index]
                            all_content.append({
                                'type': 'image',
                                'path': img_info['path'],
                                'width': img_info['width'],
                                'height': img_info['height'],
                                'ocr_text': img_info.get('ocr_text', ''),
                                'index': image_index
                            })
                            print(f"📷 이미지 {image_index + 1} 위치 확인: {img_info['original_name']}")
                            image_index += 1
                            has_image = True
                
                # 텍스트 처리
                if paragraph.text and paragraph.text.strip():
                    text = safe_korean_text(paragraph.text.strip())
                    
                    # 서식 정보 추출
                    font_size = 11  # 기본값
                    is_bold = False
                    font_name = KOREAN_FONT  # 기본 폰트 설정
                    
                    # 첫 번째 run의 서식 정보 사용
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        if hasattr(first_run, 'font'):
                            if hasattr(first_run.font, 'size') and first_run.font.size:
                                font_size = int(first_run.font.size.pt)
                            if hasattr(first_run.font, 'bold') and first_run.font.bold:
                                is_bold = True
                    
                    # 문단 스타일 확인
                    style_name = paragraph.style.name if paragraph.style else 'Normal'
                    
                    # 제목 스타일에 따른 폰트 크기 조정
                    if 'Heading 1' in style_name or 'Title' in style_name:
                        font_size = max(font_size, 18)
                    elif 'Heading 2' in style_name:
                        font_size = max(font_size, 16)
                    elif 'Heading 3' in style_name:
                        font_size = max(font_size, 14)
                    elif 'Heading 4' in style_name:
                        font_size = max(font_size, 12)
                    
                    if text:  # 텍스트가 있을 때만 추가
                        all_content.append({
                            'type': 'paragraph',
                            'content': text,
                            'font_size': font_size,
                            'is_bold': is_bold,
                            'style': style_name,
                            'font_name': font_name,
                            'index': i
                        })
                        
                        print(f"문단 {i+1}: {text[:20]}... (크기: {font_size}, 굵게: {is_bold}, 스타일: {style_name})")
                    
            except Exception as e:
                print(f"문단 {i} 처리 오류: {e}")
                continue
        
        # 3. 남은 이미지들 추가 (문단에 포함되지 않은 이미지)
        while image_index < len(extracted_images):
            img_info = extracted_images[image_index]
            all_content.append({
                'type': 'image',
                'path': img_info['path'],
                'width': img_info['width'],
                'height': img_info['height'],
                'ocr_text': img_info.get('ocr_text', ''),
                'index': image_index
            })
            print(f"📷 추가 이미지 {image_index + 1}: {img_info['original_name']}")
            image_index += 1
        
        # 4. 표 추출 (서식 포함)
        for table_idx, table in enumerate(doc.tables):
            try:
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell_idx, cell in enumerate(row.cells):
                        try:
                            cell_text = safe_korean_text(cell.text.strip())
                            if cell_text:
                                row_content.append(cell_text)
                        except:
                            row_content.append("")
                    
                    if row_content and any(row_content):
                        table_content.append(row_content)
                
                if table_content:
                    all_content.append({
                        'type': 'table',
                        'content': table_content,
                        'index': table_idx
                    })
                    print(f"표 {table_idx+1}: {len(table_content)}행")
            except Exception as e:
                print(f"표 {table_idx} 처리 오류: {e}")
                continue
        
        print(f"✅ 총 {len(all_content)}개 요소 추출 (이미지 {len(extracted_images)}개 포함)")
        return all_content
        
    except Exception as e:
        print(f"❌ DOCX 완전 서식 추출 실패: {e}")
        return []

def detect_pdf_orientation(pdf_path):
    """PDF 문서의 방향 감지"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                
                if hasattr(first_page, 'mediabox'):
                    mediabox = first_page.mediabox
                    width = float(mediabox.width)
                    height = float(mediabox.height)
                    
                    rotation = 0
                    if hasattr(first_page, 'rotation'):
                        rotation = first_page.rotation or 0
                    
                    if rotation in [90, 270]:
                        width, height = height, width
                    
                    if width > height:
                        return 'landscape', width, height
                    else:
                        return 'portrait', width, height
        
        return 'portrait', 595, 842
        
    except Exception as e:
        print(f"⚠️ PDF 방향 감지 실패: {e}")
        return 'portrait', 595, 842

def detect_docx_orientation(docx_path):
    """DOCX 문서의 방향 감지"""
    try:
        doc = Document(docx_path)
        
        if doc.sections and len(doc.sections) > 0:
            section = doc.sections[0]
            if hasattr(section, 'page_width') and hasattr(section, 'page_height'):
                width = section.page_width.inches
                height = section.page_height.inches
                
                if width > height:
                    return 'landscape'
                else:
                    return 'portrait'
        
        return 'portrait'
        
    except Exception as e:
        print(f"⚠️ DOCX 방향 감지 실패: {e}")
        return 'portrait'

def set_docx_orientation(doc, orientation):
    """DOCX 문서의 방향 설정"""
    try:
        for section in doc.sections:
            if orientation == 'landscape':
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
        return True
    except Exception as e:
        print(f"⚠️ DOCX 방향 설정 실패: {e}")
        return False

def safe_file_check(filename):
    """파일 확인 안전 확인"""
    try:
        if not filename or '.' not in filename:
            return False, 'unknown'
        
        extension = filename.lower().split('.')[-1]
        
        if extension in ['pdf', 'docx']:
            return True, extension
        else:
            return False, extension
            
    except Exception as e:
        print(f"파일 확인 오류: {e}")
        return False, 'unknown'

def clean_temp_files(file_list):
    """임시 파일 안전 삭제"""
    for file_path in file_list:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except:
            pass

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/fonts')
def list_fonts():
    """사용 가능한 폰트 목록 API"""
    font_list = []
    for font_name, font_info in AVAILABLE_FONTS.items():
        font_list.append({
            'name': font_name,
            'display_name': font_info['display_name'],
            'current': font_name == KOREAN_FONT
        })
    
    return jsonify({
        'fonts': font_list,
        'current_font': KOREAN_FONT,
        'total_fonts': len(AVAILABLE_FONTS)
    })

@app.route('/convert', methods=['POST'])
def convert_file():
    temp_files = []
    
    try:
        print("=== PDF ↔ DOCX 변환 시작 ===")
        print(f"🔤 사용 가능한 폰트: {len(AVAILABLE_FONTS)}개")
        # OCR 관련 출력 제거
        
        # 1. 파일 확인
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
        
        file = request.files['file']
        if not file.filename:
            return jsonify({'success': False, 'error': '파일명이 없습니다.'}), 400
        
        print(f"업로드된 파일: {file.filename}")
        
        # 2. 파일 형식 확인
        is_valid, extension = safe_file_check(file.filename)
        if not is_valid:
            return jsonify({
                'success': False, 
                'error': f'지원하지 않는 파일 형식입니다. PDF 또는 DOCX 파일만 업로드 가능합니다. (현재: {extension})'
            }), 400
        
        # 3. 파일 저장
        timestamp = str(int(time.time() * 1000))
        safe_filename = secure_filename(file.filename)
        name_without_ext = safe_filename.rsplit('.', 1)[0] if '.' in safe_filename else safe_filename
        input_path = os.path.join('uploads', f"{name_without_ext}_{timestamp}.{extension}")
        
        file.save(input_path)
        temp_files.append(input_path)
        print(f"✅ 파일 저장: {input_path}")
        
        # 4. 변환 처리
        if extension == 'pdf':
            # PDF → DOCX
            output_path = os.path.join('outputs', f"{name_without_ext}_{timestamp}.docx")
            
            try:
                print("📄 PDF → DOCX 변환 시작")
                
                pdf_orientation, pdf_width, pdf_height = detect_pdf_orientation(input_path)
                images = convert_from_path(input_path, dpi=150)
                
                doc = Document()
                set_docx_orientation(doc, pdf_orientation)
                
                success_count = 0
                for i, img in enumerate(images):
                    try:
                        img_path = os.path.join('uploads', f'page_{timestamp}_{i}.jpg')
                        temp_files.append(img_path)
                        
                        img.save(img_path, 'JPEG', quality=85)
                        
                        if pdf_orientation == 'landscape':
                            doc.add_picture(img_path, width=Inches(9))
                        else:
                            doc.add_picture(img_path, width=Inches(6))
                        
                        if i < len(images) - 1:
                            doc.add_page_break()
                        
                        success_count += 1
                        
                    except Exception as e:
                        print(f"⚠️ 페이지 {i+1} 처리 오류: {e}")
                        continue
                
                if success_count == 0:
                    doc.add_paragraph("PDF 변환 완료")
                    doc.add_paragraph(f"원본 파일: {file.filename}")
                
                doc.save(output_path)
                print(f"✅ DOCX 저장 완료: {success_count}개 페이지")
                
            except Exception as e:
                print(f"❌ PDF 변환 오류: {e}")
                doc = Document()
                doc.add_paragraph("PDF 변환 중 오류 발생")
                doc.save(output_path)
        
        elif extension == 'docx':
            # DOCX → PDF (다중 폰트 지원)
            output_path = os.path.join('outputs', f"{name_without_ext}_{timestamp}.pdf")
            
            try:
                print("📄 DOCX → PDF 변환 시작 (다중 폰트 지원)")
                
                # 방향 감지
                docx_orientation = detect_docx_orientation(input_path)
                
                # 강화된 서식 정보와 함께 내용 추출 (이미지 포함)
                # 557번째 줄 확인
                # 변경 전: extract_docx_with_complete_formatting_enhanced
                # 변경 후: extract_docx_with_complete_formatting
                content_list = extract_docx_with_complete_formatting(input_path, temp_files)
                
                if not content_list:
                    return jsonify({'success': False, 'error': 'DOCX 파일에서 내용을 추출할 수 없습니다.'}), 400
                
                # PDF 페이지 크기 설정
                if docx_orientation == 'landscape':
                    page_size = landscape(A4)
                    base_font_size = 10
                    line_height_base = 16
                    max_chars_per_line = 70
                else:
                    page_size = portrait(A4)
                    base_font_size = 11
                    line_height_base = 18
                    max_chars_per_line = 50
                
                # PDF 생성
                c = canvas.Canvas(output_path, pagesize=page_size)
                width, height = page_size
                
                print(f"📄 PDF 생성: {width:.0f} x {height:.0f} ({docx_orientation})")
                
                margin_left = 50
                margin_right = width - 50
                margin_top = height - 50
                margin_bottom = 50
                
                y_pos = margin_top
                
                # 내용 처리 (다중 폰트 지원)
                processed_items = 0
                image_count = 0
                
                for item in content_list:
                    try:
                        if processed_items >= 200:  # 처리 제한
                            break
                        
                        if item['type'] == 'paragraph':
                            # 문단 처리 (다중 폰트 지원)
                            text = item['content']
                            font_size = item.get('font_size', base_font_size)
                            is_bold = item.get('is_bold', False)
                            style = item.get('style', 'Normal')
                            font_name = item.get('font_name', None)
                            
                            # 줄 높이 조정 (폰트 크기에 비례)
                            line_height = max(line_height_base, int(font_size * 1.4))
                            
                            # 제목 스타일 추가 간격
                            if 'Heading' in style or 'Title' in style:
                                y_pos -= 15  # 제목 전 추가 간격
                            
                            # 줄 단위로 분할 (폰트 크기 고려)
                            chars_per_line = max(20, int(max_chars_per_line * base_font_size / font_size))
                            
                            lines = []
                            current_line = ""
                            
                            for char in text:
                                if len(current_line) >= chars_per_line:
                                    lines.append(current_line)
                                    current_line = char
                                else:
                                    current_line += char
                            
                            if current_line:
                                lines.append(current_line)
                            
                            # 각 줄 출력 (다중 폰트 적용)
                            for line in lines:
                                if y_pos < margin_bottom + line_height:
                                    c.showPage()
                                    y_pos = margin_top
                                
                                # 향상된 텍스트 그리기 (다중 폰트 적용)
                                draw_korean_text(c, margin_left, y_pos, line, font_size)
                                draw_korean_text(c, margin_left + 0.5, y_pos, line, font_size)
                                
                                y_pos -= line_height
                            
                            # 문단 간격 (스타일에 따라)
                            if 'Heading' in style or 'Title' in style:
                                y_pos -= 15  # 제목 후 추가 간격
                            else:
                                y_pos -= 8   # 일반 문단 간격
                        
                        elif item['type'] == 'image':
                            # 개선된 이미지 처리 (레이아웃 보존)
                            try:
                                image_path = item['path']
                                img_width = item['width']
                                img_height = item['height']
                                ocr_text = item.get('ocr_text', '')
                                
                                # 이미지 크기 계산 (원본 비율 유지)
                                aspect_ratio = img_width / img_height
                                max_width = (width - margin_left - 50) * 0.8  # 페이지 너비의 80%
                                max_height = (height - margin_top - margin_bottom) * 0.6  # 페이지 높이의 60%
                                
                                if max_width / aspect_ratio <= max_height:
                                    final_width = max_width
                                    final_height = max_width / aspect_ratio
                                else:
                                    final_height = max_height
                                    final_width = max_height * aspect_ratio
                                
                                # 이미지 여백 설정
                                image_margin_top = 20
                                image_margin_bottom = 15
                                
                                # 페이지 넘김 확인
                                required_space = final_height + image_margin_top + image_margin_bottom
                                if y_pos - required_space < margin_bottom:
                                    c.showPage()
                                    y_pos = margin_top
                                
                                # 이미지 중앙 정렬 계산
                                # 이미지 중앙 정렬 계산
                                page_width = width - margin_left - 50  # 사용 가능한 페이지 너비
                                image_x = margin_left + (page_width - final_width) / 2
                                
                                # 이미지 파일 존재 확인
                                if os.path.exists(image_path):
                                    try:
                                        # 메모리 최적화된 이미지 처리 (Render 환경 고려)
                                        from PIL import Image as PILImage
                                        
                                        # 이미지 크기 사전 확인 및 최적화
                                        with PILImage.open(image_path) as pil_img:
                                            # 이미지가 너무 크면 임시로 리사이즈
                                            max_dimension = 2000  # Render 환경 메모리 제한 고려
                                            if pil_img.width > max_dimension or pil_img.height > max_dimension:
                                                # 임시 리사이즈된 이미지 생성
                                                temp_img = pil_img.copy()
                                                temp_img.thumbnail((max_dimension, max_dimension), PILImage.Resampling.LANCZOS)
                                                
                                                # 임시 파일로 저장
                                                import tempfile
                                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                                                    temp_img.save(temp_file.name, 'PNG', optimize=True)
                                                    optimized_image_path = temp_file.name
                                                    temp_files.append(optimized_image_path)
                                                temp_img.close()
                                            else:
                                                optimized_image_path = image_path
                                        
                                        # 이미지 그리기 (메모리 최적화된 경로 사용)
                                        img_reader = ImageReader(optimized_image_path)
                                        c.drawImage(img_reader, image_x, y_pos - final_height, 
                                                  width=final_width, height=final_height,
                                                  preserveAspectRatio=True, anchor='c')
                                        
                                        print(f"✅ 이미지 {image_count + 1} 삽입: {final_width:.0f}x{final_height:.0f} (원본: {img_width}x{img_height})")
                                        
                                        # 이미지 하단 여백
                                        y_pos -= final_height + image_margin_bottom
                                        image_count += 1
                                        
                                        # OCR 텍스트 처리 (이미지 성공 시)
                                        if ocr_text and len(ocr_text.strip()) > 0:
                                            y_pos -= 5
                                            # OCR 텍스트를 이미지 아래 중앙에 배치
                                            ocr_lines = [ocr_text[i:i+60] for i in range(0, len(ocr_text), 60)]
                                            for ocr_line in ocr_lines[:3]:  # 최대 3줄
                                                if y_pos < margin_bottom + line_height_base:
                                                    c.showPage()
                                                    y_pos = margin_top
                                                draw_korean_text(c, margin_left + 20, y_pos, f"📝 {ocr_line}", base_font_size - 1)
                                                y_pos -= line_height_base
                                
                                    except Exception as e:
                                        error_msg = f"이미지 처리 오류 ({os.path.basename(image_path)}): {str(e)}"
                                        print(f"❌ {error_msg}")
                                        
                                        # 상세한 에러 로깅 (Render 환경 디버깅용)
                                        import traceback
                                        print(f"상세 에러 정보: {traceback.format_exc()}")
                                        
                                        # 에러 유형별 메시지 제공
                                        if "Memory" in str(e) or "memory" in str(e):
                                            error_display = f"[메모리 부족으로 이미지 로드 실패: {os.path.basename(image_path)}]"
                                        elif "Permission" in str(e) or "permission" in str(e):
                                            error_display = f"[권한 오류로 이미지 로드 실패: {os.path.basename(image_path)}]"
                                        elif "format" in str(e).lower() or "Format" in str(e):
                                            error_display = f"[지원되지 않는 이미지 형식: {os.path.basename(image_path)}]"
                                        else:
                                            error_display = f"[이미지 로드 실패: {os.path.basename(image_path)}]"
                                        
                                        # 이미지 오류 시 대체 텍스트
                                        draw_korean_text(c, margin_left, y_pos, error_display, base_font_size)
                                        y_pos -= line_height_base * 2
                                        image_count += 1
                                        
                                        # OCR 텍스트만 표시 (이미지 실패 시)
                                        if ocr_text and len(ocr_text.strip()) > 0:
                                            ocr_lines = [ocr_text[i:i+60] for i in range(0, len(ocr_text), 60)]
                                            for ocr_line in ocr_lines[:3]:  # 최대 3줄
                                                if y_pos < margin_bottom + line_height_base:
                                                    c.showPage()
                                                    y_pos = margin_top
                                                draw_korean_text(c, margin_left + 20, y_pos, f"📝 {ocr_line}", base_font_size - 1)
                                                y_pos -= line_height_base
                                
                                else:
                                    print(f"이미지 파일 없음: {image_path}")
                                    draw_korean_text(c, margin_left, y_pos, f"[이미지 {image_count + 1} - 파일 없음]", base_font_size)
                                    y_pos -= line_height_base
                                    image_count += 1
                                    
                                    # OCR 텍스트만 표시 (파일 없음 시)
                                    if ocr_text and len(ocr_text.strip()) > 0:
                                        ocr_lines = [ocr_text[i:i+60] for i in range(0, len(ocr_text), 60)]
                                        for ocr_line in ocr_lines[:3]:  # 최대 3줄
                                            if y_pos < margin_bottom + line_height_base:
                                                c.showPage()
                                                y_pos = margin_top
                                            draw_korean_text(c, margin_left + 20, y_pos, f"📝 {ocr_line}", base_font_size - 1)
                                            y_pos -= line_height_base
                                
                            except Exception as e:
                                print(f"이미지 처리 오류: {e}")
                                draw_korean_text(c, margin_left, y_pos, f"[이미지 처리 오류: {str(e)[:50]}]", base_font_size)
                                y_pos -= line_height_base * 2
                        
                        elif item['type'] == 'table':
                            # 표 처리 (다중 폰트)
                            table_data = item['content']
                            
                            # 표 제목
                            if y_pos < margin_bottom + line_height_base:
                                c.showPage()
                                y_pos = margin_top
                            
                            draw_korean_text(c, margin_left, y_pos, f"[표 {item['index'] + 1}]", base_font_size + 1, 'Heading')
                            y_pos -= line_height_base + 5
                            
                            # 표 내용
                            for row in table_data:
                                if y_pos < margin_bottom + line_height_base:
                                    c.showPage()
                                    y_pos = margin_top
                                
                                row_text = " | ".join(str(cell) for cell in row)
                                if len(row_text) > max_chars_per_line:
                                    row_text = row_text[:max_chars_per_line] + "..."
                                
                                draw_korean_text(c, margin_left + 10, y_pos, row_text, base_font_size - 1)
                                y_pos -= line_height_base
                            
                            y_pos -= 10  # 표 간격
                        
                        processed_items += 1
                        
                    except Exception as e:
                        print(f"항목 처리 오류: {e}")
                        continue
                
                c.save()
                print(f"✅ PDF 저장 완료: {processed_items}개 항목 (이미지 {image_count}개 포함)")
                
            except Exception as e:
                print(f"❌ DOCX 변환 오류: {e}")
                # 오류 시에도 기본 PDF 생성
                c = canvas.Canvas(output_path, pagesize=portrait(A4))
                draw_korean_text(c, 50, 750, "DOCX 변환 실패", 12)
                draw_korean_text(c, 50, 730, f"파일: {file.filename}", 10)
                c.save()
        
        # 5. 임시 파일 정리
        clean_temp_files(temp_files)
        
        # 6. 파일 다운로드
        if os.path.exists(output_path):
            if extension == 'pdf':
                download_name = f"{name_without_ext}.docx"
            else:
                download_name = f"{name_without_ext}.pdf"
            
            print(f"✅ 변환 완료: {download_name}")
            return send_file(output_path, as_attachment=True, download_name=download_name)
        else:
            return jsonify({'success': False, 'error': '변환된 파일을 찾을 수 없습니다.'}), 500
    
    except Exception as e:
        print(f"❌ 전체 오류: {e}")
        clean_temp_files(temp_files)
        return jsonify({'success': False, 'error': f'변환 중 오류가 발생했습니다: {str(e)}'}), 500

@app.before_request
def limit_file_size():
    if request.endpoint == 'convert_file' and request.method == 'POST':
        if 'file' in request.files:
            file = request.files['file']
            if file and len(file.read()) > 100 * 1024 * 1024:  # 100MB로 통일
                file.seek(0)  # 파일 포인터 리셋
                return jsonify({'error': '파일 크기가 너무 큽니다 (최대 100MB)'}), 413
            file.seek(0)  # 파일 포인터 리셋

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': '파일 크기가 너무 큽니다 (최대 100MB)'}), 413

@app.errorhandler(Exception)
def handle_error(e):
    return jsonify({'error': f'서버 오류: {str(e)}'}), 500

@app.route('/health')
def health():
    return "ok", 200

if __name__ == '__main__':
    print("🚀 PDF ↔ DOCX 변환기 (빠른 시작 버전)")
    print(f"🔤 한글 폰트: {KOREAN_FONT} (사용가능: {KOREAN_FONT_AVAILABLE})")
    for font_name, font_info in AVAILABLE_FONTS.items():
        print(f"   - {font_info['display_name']} ({font_name})")
    print("📄 완벽한 문서 변환")
    print("🎨 원본 서식 보존")
    print("📍 http://localhost:8080")
    app.run(debug=True, host='0.0.0.0', port=8080)