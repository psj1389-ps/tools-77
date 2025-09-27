import PyPDF2
from docx import Document
from docx.shared import Inches
from pdf2image import convert_from_path
import tempfile
import os
import logging
import re
from urllib.parse import unquote

# Local imports
from adobe_converter import AdobePDFConverter
from ocr_helper import extract_text_with_ocr

def get_safe_filename(pdf_path):
    """원본 파일명에서 안전한 파일명 추출 (확장자 제거, 특수문자 처리)"""
    try:
        # 파일명 추출
        basename = os.path.basename(pdf_path)
        # 확장자 제거
        name_without_ext = os.path.splitext(basename)[0]
        
        # URL 디코딩 (한글 파일명 처리)
        try:
            name_without_ext = unquote(name_without_ext, encoding='utf-8')
        except:
            pass
        
        # 파일명에 사용할 수 없는 문자 제거/대체
        safe_chars = re.sub(r'[<>:"/\\|?*]', '_', name_without_ext)
        # 연속된 언더스코어 정리
        safe_chars = re.sub(r'_+', '_', safe_chars)
        # 앞뒤 공백 및 언더스코어 제거
        safe_chars = safe_chars.strip('_ ')
        
        # 빈 문자열이면 기본값 사용
        if not safe_chars:
            safe_chars = 'converted_document'
            
        return safe_chars
    except Exception as e:
        logging.warning(f"파일명 처리 중 오류 발생: {e}")
        return 'converted_document'

def get_unique_filename(file_path):
    """중복 파일명 처리 - 파일이 존재하면 번호를 추가"""
    if not os.path.exists(file_path):
        return file_path
    
    directory = os.path.dirname(file_path)
    filename = os.path.basename(file_path)
    name, ext = os.path.splitext(filename)
    
    counter = 1
    while True:
        new_filename = f"{name}_{counter}{ext}"
        new_path = os.path.join(directory, new_filename)
        if not os.path.exists(new_path):
            return new_path
        counter += 1
        
        # 무한 루프 방지
        if counter > 1000:
            import time
            timestamp = int(time.time())
            new_filename = f"{name}_{timestamp}{ext}"
            return os.path.join(directory, new_filename)

def analyze_page_orientation(pdf_path):
    """PDF 페이지 방향 분석 (가로형/세로형 자동감지)"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if len(pdf_reader.pages) == 0:
                return "unknown"
            
            landscape_pages = 0
            portrait_pages = 0
            
            for page in pdf_reader.pages:
                # 페이지 크기 정보 가져오기
                mediabox = page.mediabox
                width = float(mediabox.width)
                height = float(mediabox.height)
                
                # 회전 정보 고려
                rotation = page.get('/Rotate', 0)
                if rotation in [90, 270]:
                    width, height = height, width
                
                # 가로형/세로형 판단
                if width > height:
                    landscape_pages += 1
                else:
                    portrait_pages += 1
            
            # 주요 방향 결정
            if landscape_pages > portrait_pages:
                orientation = "landscape"
                ratio = landscape_pages / (landscape_pages + portrait_pages)
            else:
                orientation = "portrait"
                ratio = portrait_pages / (landscape_pages + portrait_pages)
            
            logging.info(f"페이지 방향 분석: {orientation} ({ratio:.2f} 비율)")
            return {
                "orientation": orientation,
                "ratio": ratio,
                "landscape_pages": landscape_pages,
                "portrait_pages": portrait_pages
            }
            
    except Exception as e:
        logging.error(f"페이지 방향 분석 중 오류 발생: {e}")
        return {"orientation": "unknown", "ratio": 0}

def detect_official_document(pdf_path):
    """공문서 자동감지 (텍스트 패턴, 레이아웃 분석)"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if len(pdf_reader.pages) == 0:
                return False
            
            # 첫 페이지에서 텍스트 추출
            first_page_text = pdf_reader.pages[0].extract_text()
            
            # 공문서 키워드 패턴
            official_keywords = [
                '공문', '시행', '수신', '발신', '시행일자', '문서번호', '담당부서',
                '결재', '시장', '구청장', '과장', '팀장', '담당자',
                '붙임', '끝.', '협조사항', '시행근거', '추진계획',
                '○', '가.', '나.', '다.', '라.', '마.',
                '1.', '2.', '3.', '4.', '5.',
                '기안자', '검토자', '결재권자', '시행자'
            ]
            
            # 공문서 레이아웃 패턴
            layout_patterns = [
                r'문서번호\s*:', r'시행일자\s*:', r'수신\s*:', r'발신\s*:',
                r'제\s*목\s*:', r'담당부서\s*:', r'담당자\s*:',
                r'\d{4}-\d+', r'\d{4}\.\d{1,2}\.\d{1,2}',  # 문서번호, 날짜 패턴
                r'붙임\s*\d*\s*부', r'끝\s*\.',
                r'[가-힣]+시장|[가-힣]+구청장|[가-힣]+과장'
            ]
            
            keyword_count = 0
            pattern_count = 0
            
            # 키워드 검사
            for keyword in official_keywords:
                if keyword in first_page_text:
                    keyword_count += 1
            
            # 패턴 검사
            for pattern in layout_patterns:
                if re.search(pattern, first_page_text):
                    pattern_count += 1
            
            # 공문서 판단 기준
            is_official = (keyword_count >= 3 or pattern_count >= 2)
            confidence = (keyword_count * 0.3 + pattern_count * 0.7) / 10
            
            logging.info(f"공문서 감지: {'예' if is_official else '아니오'} (키워드: {keyword_count}, 패턴: {pattern_count}, 신뢰도: {confidence:.2f})")
            
            return {
                "is_official": is_official,
                "confidence": min(confidence, 1.0),
                "keyword_count": keyword_count,
                "pattern_count": pattern_count
            }
            
    except Exception as e:
        logging.error(f"공문서 감지 중 오류 발생: {e}")
        return {"is_official": False, "confidence": 0}

def analyze_pdf_content(pdf_path):
    """PDF 내용 분석하여 타입 결정 (방향 및 공문서 정보 포함)"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            if total_pages == 0:
                return {"type": "empty"}

            text_pages = 0
            for page in pdf_reader.pages:
                # 페이지에서 텍스트 추출 시도
                text = page.extract_text()
                if text and len(text.strip()) > 50:  # 50자 이상이면 텍스트 페이지로 간주
                    text_pages += 1
            
            text_ratio = text_pages / total_pages
            
            # 페이지 방향 분석
            orientation_info = analyze_page_orientation(pdf_path)
            
            # 공문서 감지
            official_info = detect_official_document(pdf_path)
            
            # PDF 타입 결정
            if text_ratio > 0.8:
                pdf_type = "text_based"
                logging.info(f"PDF 분석 결과: 텍스트 기반 (텍스트 비율: {text_ratio:.2f})")
            elif text_ratio < 0.2:
                pdf_type = "scanned_image"
                logging.info(f"PDF 분석 결과: 이미지 기반 (텍스트 비율: {text_ratio:.2f})")
            else:
                pdf_type = "mixed"
                logging.info(f"PDF 분석 결과: 혼합형 (텍스트 비율: {text_ratio:.2f})")
            
            return {
                "type": pdf_type,
                "text_ratio": text_ratio,
                "orientation": orientation_info,
                "official_document": official_info
            }
            
    except Exception as e:
        logging.error(f"PDF 내용 분석 중 오류 발생: {e}")
        return {"type": "unknown"}

def convert_image_pdf_to_docx(pdf_path, use_ocr=False):
    """이미지 기반 PDF를 DOCX로 변환 (OCR 옵션 포함)"""
    # 원본 파일명을 유지하여 출력 파일명 생성
    original_name = get_safe_filename(pdf_path)
    suffix = '_ocr' if use_ocr else '_image'
    filename = f"{original_name}{suffix}.docx"
    outputs_dir = os.path.join(os.path.dirname(pdf_path), '..', 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    output_path = get_unique_filename(os.path.join(outputs_dir, filename))
    try:
        images = convert_from_path(pdf_path, dpi=200)
        doc = Document()

        if use_ocr:
            logging.info("OCR을 사용하여 이미지에서 텍스트를 추출합니다.")
            texts = extract_text_with_ocr(pdf_path)
            for text in texts:
                doc.add_paragraph(text)
                doc.add_page_break()
        else:
            logging.info("이미지를 원본 그대로 DOCX에 삽입합니다.")
            for i, image in enumerate(images):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_image:
                    image.save(temp_image.name, 'JPEG')
                    doc.add_picture(temp_image.name, width=Inches(6.0))
                    if i < len(images) - 1:
                        doc.add_page_break()
                os.unlink(temp_image.name)
        
        doc.save(output_path)
        logging.info(f"이미지 기반 PDF를 DOCX로 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"이미지 기반 PDF 변환 오류: {e}")
        return None

def fallback_text_conversion(pdf_path):
    """PyPDF2를 사용한 텍스트 추출 (폴백)"""
    # 원본 파일명을 유지하여 출력 파일명 생성
    original_name = get_safe_filename(pdf_path)
    filename = f"{original_name}_text.docx"
    outputs_dir = os.path.join(os.path.dirname(pdf_path), '..', 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    output_path = get_unique_filename(os.path.join(outputs_dir, filename))
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            doc = Document()
            for page in pdf_reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(output_path)
        logging.info(f"폴백 텍스트 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"폴백 텍스트 변환 오류: {e}")
        return None

def convert_official_document(pdf_path, analysis_result):
    """공문서 특화 변환 - 레이아웃 보존 최우선"""
    try:
        # 공문서는 레이아웃 보존이 중요하므로 Adobe API 우선 사용
        adobe_converter = AdobePDFConverter()
        if adobe_converter.execution_context:
            logging.info("공문서 Adobe API 변환 시도")
            result = adobe_converter.convert_to_docx_official(pdf_path, analysis_result)
            if result:
                return result
        
        # Adobe API 실패 시 고품질 이미지 변환
        logging.info("공문서 고품질 이미지 변환으로 폴백")
        return convert_image_pdf_to_docx_official(pdf_path, analysis_result)
        
    except Exception as e:
        logging.error(f"공문서 변환 중 오류: {e}")
        return convert_image_pdf_to_docx(pdf_path, use_ocr=False)

def convert_image_pdf_to_docx_optimized(pdf_path, analysis_result, use_ocr=False):
    """방향별 최적화가 적용된 이미지 PDF 변환"""
    orientation_info = analysis_result.get("orientation", {})
    orientation = orientation_info.get("orientation", "portrait")
    
    # 원본 파일명을 유지하여 출력 파일명 생성
    original_name = get_safe_filename(pdf_path)
    suffix = f'_{orientation}_ocr' if use_ocr else f'_{orientation}_image'
    filename = f"{original_name}{suffix}.docx"
    outputs_dir = os.path.join(os.path.dirname(pdf_path), '..', 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    output_path = get_unique_filename(os.path.join(outputs_dir, filename))
    
    try:
        # 방향에 따른 DPI 최적화
        dpi = 300 if orientation == "landscape" else 200
        images = convert_from_path(pdf_path, dpi=dpi)
        doc = Document()
        
        # 방향에 따른 페이지 설정
        if orientation == "landscape":
            from docx.enum.section import WD_ORIENT
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        if use_ocr:
            logging.info(f"OCR을 사용하여 {orientation} 이미지에서 텍스트를 추출합니다.")
            texts = extract_text_with_ocr(pdf_path)
            for text in texts:
                doc.add_paragraph(text)
                doc.add_page_break()
        else:
            logging.info(f"{orientation} 이미지를 원본 그대로 DOCX에 삽입합니다.")
            for i, image in enumerate(images):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_image:
                    image.save(temp_image.name, 'JPEG')
                    # 방향에 따른 이미지 크기 조정
                    width = Inches(8.0) if orientation == "landscape" else Inches(6.0)
                    doc.add_picture(temp_image.name, width=width)
                    if i < len(images) - 1:
                        doc.add_page_break()
                os.unlink(temp_image.name)
        
        doc.save(output_path)
        logging.info(f"{orientation} 최적화된 이미지 PDF를 DOCX로 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"최적화된 이미지 PDF 변환 오류: {e}")
        return convert_image_pdf_to_docx(pdf_path, use_ocr=use_ocr)

def convert_image_pdf_to_docx_official(pdf_path, analysis_result):
    """공문서용 고품질 이미지 변환"""
    # 원본 파일명을 유지하여 출력 파일명 생성
    original_name = get_safe_filename(pdf_path)
    filename = f"{original_name}_공문서.docx"
    outputs_dir = os.path.join(os.path.dirname(pdf_path), '..', 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    output_path = get_unique_filename(os.path.join(outputs_dir, filename))
    
    try:
        # 공문서는 고해상도로 변환
        images = convert_from_path(pdf_path, dpi=400)
        doc = Document()
        
        logging.info("공문서 고품질 이미지를 DOCX에 삽입합니다.")
        for i, image in enumerate(images):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_image:
                image.save(temp_image.name, 'JPEG', quality=95)
                doc.add_picture(temp_image.name, width=Inches(7.5))
                if i < len(images) - 1:
                    doc.add_page_break()
            os.unlink(temp_image.name)
        
        doc.save(output_path)
        logging.info(f"공문서 고품질 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"공문서 이미지 변환 오류: {e}")
        return convert_image_pdf_to_docx(pdf_path, use_ocr=False)

def fallback_text_conversion_optimized(pdf_path, analysis_result):
    """방향별 최적화가 적용된 텍스트 추출 (폴백)"""
    orientation_info = analysis_result.get("orientation", {})
    orientation = orientation_info.get("orientation", "portrait")
    
    # 원본 파일명을 유지하여 출력 파일명 생성
    original_name = get_safe_filename(pdf_path)
    filename = f"{original_name}_{orientation}_text.docx"
    outputs_dir = os.path.join(os.path.dirname(pdf_path), '..', 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    output_path = get_unique_filename(os.path.join(outputs_dir, filename))
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            doc = Document()
            
            # 방향에 따른 페이지 설정
            if orientation == "landscape":
                from docx.enum.section import WD_ORIENT
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width
            
            for page in pdf_reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(output_path)
        logging.info(f"{orientation} 최적화된 폴백 텍스트 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"최적화된 폴백 텍스트 변환 오류: {e}")
        return fallback_text_conversion(pdf_path)

def hybrid_conversion_optimized(pdf_path, analysis_result):
    """방향별 최적화가 적용된 혼합형 PDF 처리"""
    orientation_info = analysis_result.get("orientation", {})
    orientation = orientation_info.get("orientation", "portrait")
    
    logging.info(f"혼합형 PDF를 {orientation} 최적화로 이미지 기반 처리합니다.")
    return convert_image_pdf_to_docx_optimized(pdf_path, analysis_result, use_ocr=False)

def hybrid_conversion(pdf_path):
    """혼합형 PDF 처리 (현재는 이미지 기반으로 처리)"""
    logging.info("혼합형 PDF는 현재 이미지 기반으로 처리됩니다.")
    return convert_image_pdf_to_docx(pdf_path, use_ocr=False)

def smart_pdf_to_docx(pdf_path, options=None):
    """지능형 PDF 변환 - 문서 타입, 방향, 공문서 자동 감지"""
    if options is None:
        options = {}

    # 1. PDF 종합 분석
    analysis_result = analyze_pdf_content(pdf_path)
    pdf_type = analysis_result.get("type", "unknown")
    orientation_info = analysis_result.get("orientation", {})
    official_info = analysis_result.get("official_document", {})
    
    # 2. 분석 결과 로깅
    orientation = orientation_info.get("orientation", "unknown")
    is_official = official_info.get("is_official", False)
    
    logging.info(f"📊 PDF 종합 분석 완료:")
    logging.info(f"   - 문서 타입: {pdf_type}")
    logging.info(f"   - 페이지 방향: {orientation}")
    logging.info(f"   - 공문서 여부: {'예' if is_official else '아니오'}")
    
    # 3. 공문서 특화 처리
    if is_official and official_info.get("confidence", 0) > 0.5:
        logging.info("📋 공문서 감지 - 레이아웃 보존 모드로 변환")
        return convert_official_document(pdf_path, analysis_result)
    
    # 4. 일반 문서 처리 (방향별 최적화 적용)
    if pdf_type == "text_based":
        logging.info("📄 텍스트 기반 PDF 감지 - Adobe API 사용 시도")
        adobe_converter = AdobePDFConverter()
        if adobe_converter.execution_context:
            result = adobe_converter.convert_to_docx_optimized(pdf_path, analysis_result)
            if result:
                return result
        
        logging.warning("Adobe API 사용 불가 또는 실패. 폴백 텍스트 추출을 시도합니다.")
        return fallback_text_conversion_optimized(pdf_path, analysis_result)
    
    elif pdf_type == "scanned_image":
        logging.info("🖼️ 이미지 기반 PDF 감지 - 방향별 최적화 적용")
        use_ocr = options.get('force_ocr', False)
        return convert_image_pdf_to_docx_optimized(pdf_path, analysis_result, use_ocr=use_ocr)
    
    elif pdf_type == "mixed":
        logging.info("🔄 혼합형 PDF 감지 - 하이브리드 처리")
        return hybrid_conversion_optimized(pdf_path, analysis_result)
        
    else:
        logging.error("알 수 없는 PDF 타입입니다. 기본 변환을 시도합니다.")
        return convert_image_pdf_to_docx(pdf_path)